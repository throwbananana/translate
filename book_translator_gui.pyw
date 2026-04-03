#! python
# -*- coding: utf-8 -*-
"""
书籍翻译工具 GUI v2.3.1
支持PDF、TXT、EPUB、DOCX、Markdown格式的书籍翻译
可接入Gemini API、OpenAI API、Claude API、DeepSeek API等多种翻译API
支持动态添加多个本地模型，翻译与解析功能可独立选择不同API
云端配额耗尽时可自动切换到本地模型
新增：翻译记忆库（避免重复翻译）、术语表管理（统一专业术语）
v2.3 新增：PDF OCR扫描件支持 (pdfplumber+pdf2image)、可视化术语表编辑、章节目录导航、批量任务断点续传

依赖库:
py -m pip install PyPDF2 pdfplumber pdf2image ebooklib beautifulsoup4 google-generativeai openai requests python-docx
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext, simpledialog
import threading
import sys
import os
import time
import datetime
from pathlib import Path
import re
import json
import shutil
from copy import deepcopy
import hmac
import hashlib
import uuid
from concurrent.futures import ThreadPoolExecutor, as_completed

# 文件读取相关
try:
    import PyPDF2
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    import ebooklib
    from ebooklib import epub
    from bs4 import BeautifulSoup
    EPUB_SUPPORT = True
except ImportError:
    EPUB_SUPPORT = False

# API相关
try:
    import google.generativeai as genai
    GEMINI_SUPPORT = True
except ImportError:
    GEMINI_SUPPORT = False

try:
    import openai
    OPENAI_SUPPORT = True
except ImportError:
    OPENAI_SUPPORT = False

try:
    import anthropic
    CLAUDE_SUPPORT = True
except ImportError:
    CLAUDE_SUPPORT = False

try:
    import requests
    REQUESTS_SUPPORT = True
except ImportError:
    REQUESTS_SUPPORT = False

from file_processor import FileProcessor
from app_paths import get_runtime_file
from config_manager import ConfigManager, get_config_manager
from translation_memory import TranslationMemory, get_translation_memory
from glossary_manager import GlossaryManager, get_glossary_manager
from online_search import OnlineSearchManager
from translation_engine import (
    TranslationEngine,
    APIConfig,
    APIProvider,
    provider_enum_for_name,
)
from cost_estimator import CostEstimator
from docx_handler import DocxHandler
from audio_manager import AudioManager
from smart_glossary import SmartGlossaryExtractor
from book_hunter import BookHunter
from web_importer import WebImporter
from tm_editor import TMEditorDialog
from format_converter import FormatConverterDialog
from cloud_upload import CloudUploader
from community_manager import CommunityManager
from provider_utils import (
    choose_fallback_provider,
    provider_error_message,
    provider_ready,
    validate_builtin_provider,
)
from translation_review import (
    apply_manual_translation as apply_manual_translation_review,
    is_translation_incomplete as review_is_translation_incomplete,
    verify_and_retry_segments as review_verify_and_retry_segments,
)
from ui.analysis_panel import AnalysisPanel
from ui.content_notebook import ContentNotebook
from ui.failed_segments_panel import FailedSegmentsPanel
from ui.glossary_dialog import GlossaryEditorDialog
from ui.library_panel import LibraryPanel
from ui.toc_panel import TocPanel
from ui.workstation import ActionBar, ApiPanel, FilePanel, ProgressPanel

DEFAULT_TARGET_LANGUAGE = "中文"
DEFAULT_LM_STUDIO_CONFIG = {
    'api_key': 'lm-studio',
    'model': 'qwen2.5-7b-instruct-1m',
    'base_url': 'http://127.0.0.1:1234/v1'
}

DEFAULT_API_CONFIGS = {
    'gemini': {'api_key': '', 'model': 'gemini-2.5-flash'},
    'openai': {'api_key': '', 'model': 'gpt-3.5-turbo', 'base_url': ''},
    'claude': {'api_key': '', 'model': 'claude-3-haiku-20240307'},
    'deepseek': {'api_key': '', 'model': 'deepseek-chat', 'base_url': 'https://api.deepseek.com/v1'},
    'custom': {'api_key': '', 'model': '', 'base_url': ''},
    'lm_studio': deepcopy(DEFAULT_LM_STUDIO_CONFIG)
}

# 应用版本号
APP_VERSION = "2.3.1"

# 配置文件版本号
CONFIG_VERSION = APP_VERSION

class BookTranslatorGUI:
    """书籍翻译工具主界面"""

    def __init__(self, root):
        self.root = root
        self.root.title(f"书籍翻译工具 v{APP_VERSION} - 翻译记忆+术语表+多本地模型")
        self.root.geometry("950x750")

        # 初始化辅助模块
        self.file_processor = FileProcessor()
        self.web_importer = WebImporter()
        self.progress_cache_path = get_runtime_file('translation_cache.json')
        self.batch_queue_path = get_runtime_file('batch_tasks.json')

        # 初始化新模块
        self.config_manager = get_config_manager()
        self.translation_memory = get_translation_memory()
        self.glossary_manager = get_glossary_manager()
        self.online_search_manager = OnlineSearchManager(self.config_manager)
        self.community_manager = CommunityManager()
        
        # 初始化翻译引擎
        self.translation_engine = TranslationEngine()
        self.translation_engine.set_translation_memory(self.translation_memory)
        self.translation_engine.set_glossary_manager(self.glossary_manager)

        # 初始化拓展模块
        self.audio_manager = AudioManager()
        self.docx_handler = None  # 仅在加载 DOCX 时初始化
        self.smart_glossary = SmartGlossaryExtractor(self.translation_engine)
        self.book_hunter = BookHunter(self.translation_engine, self.online_search_manager)
        self.current_theme = "light"

        # 程序退出时自动保存配置
        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)

        # 翻译状态
        self.is_translating = False
        self.current_text = ""
        self.translated_text = ""
        self.translation_thread = None
        self.source_segments = []
        self.translated_segments = []
        self.failed_segments = []
        self.selected_failed_index = None
        # 是否已启用本地LM Studio备用方案
        self.lm_studio_fallback_active = False
        # 进度缓存/恢复控制
        self.text_signature = None
        self.resume_from_index = 0
        self.max_consecutive_failures = 3
        self.consecutive_failures = 0
        self.paused_due_to_failures = False

        # 大文件处理
        self.show_full_text = False
        self.preview_limit = 10000  # 预览显示前10000字符

        # 批量处理状态
        self.batch_queue = []
        self.load_batch_queue() # Load persistence
        self.is_batch_mode = False
        self.batch_window = None
        self.batch_output_dir = ""

        # 双栏对照状态
        self.sync_scroll_enabled = True

        # API配置
        self.api_configs = deepcopy(DEFAULT_API_CONFIGS)
        self.custom_local_models = {}  # 自定义本地模型存储
        self.target_language_var = tk.StringVar(value=DEFAULT_TARGET_LANGUAGE)

        # 独立的翻译和解析API选择
        self.translation_api_var = tk.StringVar(value="Gemini API")
        self.analysis_api_var = tk.StringVar(value="Gemini API")
        retry_default_api = "本地 LM Studio" if OPENAI_SUPPORT else "Gemini API"
        self.retry_api_var = tk.StringVar(value=retry_default_api)

        # 解析相关状态
        self.analysis_segments = []  # 每段的解析结果
        self.is_analyzing = False
        self.analysis_thread = None

        self.setup_ui()
        self.load_config()
        self.try_resume_cached_progress()

    def load_batch_queue(self):
        """加载批量任务列表"""
        try:
            path = self.batch_queue_path
            if path.exists():
                with open(path, 'r', encoding='utf-8') as f:
                    self.batch_queue = json.load(f)
        except Exception as e:
            print(f"Failed to load batch queue: {e}")
            self.batch_queue = []

    def save_batch_queue(self):
        """保存批量任务列表"""
        try:
            self.batch_queue_path.parent.mkdir(parents=True, exist_ok=True)
            with open(self.batch_queue_path, 'w', encoding='utf-8') as f:
                json.dump(self.batch_queue, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"Failed to save batch queue: {e}")

    def _get_support_formats_text(self):
        """构建文件选择区支持格式提示文本。"""
        format_labels = {
            ".txt": "TXT",
            ".md": "Markdown",
            ".markdown": "Markdown",
            ".pdf": "PDF",
            ".epub": "EPUB",
            ".docx": "DOCX",
            ".rtf": "RTF",
        }
        formats = []
        for ext in self.file_processor.get_supported_formats():
            label = format_labels.get(ext)
            if label and label not in formats:
                formats.append(label)
        return f"支持格式: {', '.join(formats)}"

    def _builtin_provider_enum(self, api_type):
        """将内置 provider 名称映射为枚举。"""
        return provider_enum_for_name(api_type) or APIProvider.GEMINI

    def _engine_provider_name(self, api_type):
        """翻译引擎使用的 provider 标识。"""
        provider_enum = provider_enum_for_name(api_type)
        if provider_enum is not None:
            return provider_enum.value
        if api_type in self.custom_local_models:
            return api_type
        return APIProvider.GEMINI.value

    def setup_ui(self):
        """设置用户界面"""
        # 创建菜单栏
        self.create_menu_bar()

        # 创建主框架 (Root container)
        root_container = ttk.Frame(self.root, padding="5")
        root_container.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        root_container.columnconfigure(0, weight=1)
        root_container.rowconfigure(0, weight=1)

        # === 顶部主分页 (Main Notebook) ===
        self.main_notebook = ttk.Notebook(root_container)
        self.main_notebook.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # TAB 1: 翻译工作台
        self.workstation_frame = ttk.Frame(self.main_notebook, padding="10")
        self.main_notebook.add(self.workstation_frame, text="翻译工作台")
        
        # TAB 2: 在线书城 (大页签)
        self.library_frame = ttk.Frame(self.main_notebook, padding="10")
        self.main_notebook.add(self.library_frame, text="在线书城")
        
        # 初始化搜索页签
        self.setup_search_tab()

        # 配置工作台权重
        self.workstation_frame.columnconfigure(0, weight=1)
        self.workstation_frame.rowconfigure(2, weight=1) # content_frame is row 2
        
        # 指向工作台，保持后续代码兼容
        main_frame = self.workstation_frame
        api_names = [name for name, _, _ in self.get_all_available_apis()]

        # 1. 文件选择区域
        self.file_panel = FilePanel(
            main_frame,
            support_formats_text=self._get_support_formats_text(),
            on_browse_file=self.browse_file,
            on_open_batch_window=self.open_batch_window,
            on_open_glossary_editor=self.open_glossary_editor,
            on_toggle_preview=self.toggle_full_text_display,
        )
        self.file_path_var = self.file_panel.file_path_var
        self.file_info_var = self.file_panel.file_info_var
        self.cost_var = self.file_panel.cost_var
        self.toggle_preview_btn = self.file_panel.toggle_preview_btn

        # 2. API配置区域（重构：双下拉框 + 本地模型管理）
        self.api_panel = ApiPanel(
            main_frame,
            api_names=api_names,
            translation_api_var=self.translation_api_var,
            analysis_api_var=self.analysis_api_var,
            target_language_var=self.target_language_var,
            config_manager=self.config_manager,
            on_api_type_change=self.on_api_type_change,
            on_open_translation_config=lambda: self.open_api_config_for('translation'),
            on_open_analysis_config=lambda: self.open_api_config_for('analysis'),
            on_add_local_model=self.open_add_local_model_dialog,
            on_manage_local_models=self.open_manage_local_models_dialog,
            on_update_concurrency_label=self.update_concurrency_label,
        )
        self.translation_api_combo = self.api_panel.translation_api_combo
        self.analysis_api_combo = self.api_panel.analysis_api_combo
        self.api_status_var = self.api_panel.api_status_var
        self.api_status_label = self.api_panel.api_status_label
        self.style_var = self.api_panel.style_var
        self.concurrency_var = self.api_panel.concurrency_var
        self.concurrency_scale = self.api_panel.concurrency_scale
        self.concurrency_label = self.api_panel.concurrency_label
        self.concurrency_hint_var = self.api_panel.concurrency_hint_var
        self.update_concurrency_label(self.concurrency_var.get())

        # 兼容旧代码：保留api_type_var映射
        self.api_type_var = self.translation_api_var

        # 3. 翻译内容显示区域
        content_frame = ttk.LabelFrame(main_frame, text="翻译内容", padding="10")
        content_frame.grid(row=2, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(0, 10))
        content_frame.columnconfigure(0, weight=1)
        content_frame.rowconfigure(0, weight=1)

        # === 侧边栏与主内容分割 ===
        self.content_paned = tk.PanedWindow(content_frame, orient=tk.HORIZONTAL, sashrelief=tk.RAISED)
        self.content_paned.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))

        # 左侧目录树 (TOC)
        self.toc_panel = TocPanel(self.content_paned, on_toc_click=self.on_toc_click)
        self.sidebar_frame = self.toc_panel.frame
        self.toc_tree = self.toc_panel.toc_tree

        # 右侧主要 Notebook
        self.sync_scroll_var = tk.BooleanVar(value=True)
        self.content_notebook = ContentNotebook(
            self.content_paned,
            sync_scroll_var=self.sync_scroll_var,
            on_save_comparison_edits=self.save_comparison_edits,
            on_source_scroll=self._on_source_scroll,
            on_target_scroll=self._on_target_scroll,
            on_mousewheel=self._on_mousewheel,
        )
        self.notebook = self.content_notebook.notebook
        self.original_text = self.content_notebook.original_text
        self.translated_text_widget = self.content_notebook.translated_text_widget
        self.comp_source_text = self.content_notebook.comp_source_text
        self.comp_target_text = self.content_notebook.comp_target_text

        if self.retry_api_var.get() not in api_names and api_names:
            preferred_retry = "本地 LM Studio" if "本地 LM Studio" in api_names else self.translation_api_var.get()
            if preferred_retry not in api_names:
                preferred_retry = api_names[0]
            self.retry_api_var.set(preferred_retry)

        self.failed_status_var = tk.StringVar(value="暂无失败段落")
        self.failed_panel = FailedSegmentsPanel(
            self.notebook,
            api_names=api_names,
            retry_api_var=self.retry_api_var,
            failed_status_var=self.failed_status_var,
            on_failed_select=self.on_failed_select,
            on_open_retry_config=lambda: self.open_api_config_for('retry'),
            on_retry_translation=self.retry_failed_segment,
            on_save_manual_translation=self.save_manual_translation,
        )
        self.failed_listbox = self.failed_panel.failed_listbox
        self.failed_source_text = self.failed_panel.failed_source_text
        self.manual_translation_text = self.failed_panel.manual_translation_text
        self.retry_api_combo = self.failed_panel.retry_api_combo

        self.analysis_status_var = tk.StringVar(value="翻译完成后可进行解析")
        self.analysis_panel = AnalysisPanel(
            self.notebook,
            analysis_status_var=self.analysis_status_var,
            on_analysis_segment_select=self.on_analysis_segment_select,
            on_analyze_selected_segment=self.analyze_selected_segment,
            on_copy_analysis_content=self.copy_analysis_content,
        )
        self.analysis_listbox = self.analysis_panel.analysis_listbox
        self.analysis_text = self.analysis_panel.analysis_text

        # 4. 进度和控制区域
        self.progress_panel = ProgressPanel(main_frame)
        self.progress_var = self.progress_panel.progress_var
        self.progress_bar = self.progress_panel.progress_bar
        self.progress_text_var = self.progress_panel.progress_text_var

        # 5. 操作按钮区域（分两行）
        self.action_bar = ActionBar(
            main_frame,
            on_start_translation=self.start_translation,
            on_stop_translation=self.stop_translation,
            on_start_batch_analysis=self.start_batch_analysis,
            on_stop_analysis=self.stop_analysis,
            on_export_translation=self.export_translation,
            on_export_bilingual_epub=self.export_bilingual_epub,
            on_export_audiobook=self.export_audiobook,
            on_export_analysis=self.export_analysis,
            on_clear_all=self.clear_all,
        )
        self.translate_btn = self.action_bar.translate_btn
        self.stop_btn = self.action_bar.stop_btn
        self.analyze_all_btn = self.action_bar.analyze_all_btn
        self.stop_analysis_btn = self.action_bar.stop_analysis_btn

    def _on_source_scroll(self, *args):
        self.comp_source_text.yview(*args)
        if self.sync_scroll_var.get():
            self.comp_target_text.yview_moveto(args[0])

    def _on_target_scroll(self, *args):
        self.comp_target_text.yview(*args)
        if self.sync_scroll_var.get():
            self.comp_source_text.yview_moveto(args[0])
            
    def _on_mousewheel(self, event, other_widget):
        if self.sync_scroll_var.get():
            # 传递滚轮事件给另一个控件
            other_widget.yview_scroll(int(-1*(event.delta/120)), "units")
            
    def update_comparison_view(self):
        """更新双栏对照视图的内容"""
        # 准备原文文本（按段落分隔）
        source_display = "\n\n".join(self.source_segments) if self.source_segments else self.current_text
        
        # 准备译文文本（确保与原文段落对应）
        target_segments_display = list(self.translated_segments)
        # 补齐长度
        if len(target_segments_display) < len(self.source_segments):
            target_segments_display.extend([""] * (len(self.source_segments) - len(target_segments_display)))
            
        target_display = "\n\n".join(target_segments_display)
        
        self.comp_source_text.config(state='normal')
        self.comp_source_text.delete('1.0', tk.END)
        self.comp_source_text.insert('1.0', source_display)
        self.comp_source_text.config(state='disabled')
        
        self.comp_target_text.delete('1.0', tk.END)
        self.comp_target_text.insert('1.0', target_display)

    def save_comparison_edits(self):
        """保存双栏视图中的修改到主数据并同步到记忆库"""
        new_text = self.comp_target_text.get('1.0', tk.END).strip()
        if not new_text:
            return
            
        # 尝试按双换行符分割回段落
        new_segments = re.split(r'\n\s*\n', new_text)
        
        # 简单的完整性检查
        if abs(len(new_segments) - len(self.source_segments)) > 5:
            confirm = messagebox.askyesno(
                "段落数量不匹配", 
                f"编辑后的段落数 ({len(new_segments)}) 与原文段落数 ({len(self.source_segments)}) 差异较大。\n"
                "这可能导致后续对照错位。\n\n是否仍要保存？"
            )
            if not confirm:
                return

        # 同步到翻译记忆库 (Linkage 1)
        count_updated = 0
        target_lang = self.get_target_language()
        
        # 比较差异并保存
        limit = min(len(new_segments), len(self.source_segments), len(self.translated_segments))
        for i in range(limit):
            old_trans = self.translated_segments[i]
            new_trans = new_segments[i]
            source = self.source_segments[i]
            
            # 如果有实质性修改
            if old_trans.strip() != new_trans.strip() and source.strip():
                try:
                    self.translation_memory.store(
                        source_text=source,
                        translated_text=new_trans,
                        target_lang=target_lang,
                        api_provider="manual",
                        model="user_correction",
                        quality_score=100
                    )
                    count_updated += 1
                except Exception as e:
                    print(f"Sync to TM failed for segment {i}: {e}")

        # 更新主数据
        self.translated_segments = new_segments
        self.rebuild_translated_text()
        self.save_progress_cache()
        
        msg = "修改已保存并同步到主视图"
        if count_updated > 0:
            msg += f"\n\n已将 {count_updated} 处人工修正同步到翻译记忆库！"
            
        messagebox.showinfo("成功", msg)

    def open_glossary_editor(self):
        """打开术语表编辑器"""
        GlossaryEditorDialog(self.root, self.glossary_manager)

    def open_cloud_share(self):
        """打开云端分享对话框"""
        dialog = tk.Toplevel(self.root)
        dialog.title("云端分享 (Cloud Share)")
        dialog.geometry("600x480")
        dialog.transient(self.root)
        dialog.grab_set()
        
        frame = ttk.Frame(dialog, padding=20)
        frame.pack(fill=tk.BOTH, expand=True)
        
        # 1. 文件选择
        ttk.Label(frame, text="选择要上传的文件:", font=("", 10, "bold")).pack(anchor=tk.W, pady=(0, 5))
        
        file_frame = ttk.Frame(frame)
        file_frame.pack(fill=tk.X, pady=5)
        
        file_path_var = tk.StringVar()
        if self.file_path_var.get():
            file_path_var.set(self.file_path_var.get())
            
        entry = ttk.Entry(file_frame, textvariable=file_path_var)
        entry.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        def browse():
            f = filedialog.askopenfilename()
            if f: file_path_var.set(f)
            
        ttk.Button(file_frame, text="浏览...", command=browse).pack(side=tk.LEFT, padx=5)
        
        # 快捷选项
        quick_frame = ttk.Frame(frame)
        quick_frame.pack(fill=tk.X, pady=5)
        
        def set_current_txt():
            if not self.translated_text:
                messagebox.showinfo("提示", "当前没有译文")
                return
            try:
                # Save to a temporary file
                temp_dir = Path("temp_exports")
                temp_dir.mkdir(exist_ok=True)
                
                # Use original filename + _translated.txt
                orig_name = Path(self.file_path_var.get()).stem if self.file_path_var.get() else "translation"
                temp_path = temp_dir / f"{orig_name}_translated.txt"
                
                with open(temp_path, 'w', encoding='utf-8') as f:
                    f.write(self.translated_text)
                file_path_var.set(str(temp_path.absolute()))
            except Exception as e:
                messagebox.showerror("错误", str(e))
                
        ttk.Button(quick_frame, text="当前译文 (TXT)", command=set_current_txt).pack(side=tk.LEFT, padx=2)
        
        # 2. 服务选择
        ttk.Label(frame, text="选择分享服务:", font=("", 10, "bold")).pack(anchor=tk.W, pady=(15, 5))
        service_var = tk.StringVar(value="Catbox (永久/长期)")
        
        s1 = ttk.Radiobutton(frame, text="Catbox (推荐 - 永久有效，最大200MB)", variable=service_var, value="Catbox (永久/长期)")
        s1.pack(anchor=tk.W)
        
        s2 = ttk.Radiobutton(frame, text="File.io (一次性 - 下载1次或2周后删除)", variable=service_var, value="File.io (一次性/2周)")
        s2.pack(anchor=tk.W)
        
        s3 = ttk.Radiobutton(frame, text="Litterbox (临时 - 72小时后删除)", variable=service_var, value="Litterbox (72小时)")
        s3.pack(anchor=tk.W)
            
        # 3. 上传按钮
        status_var = tk.StringVar(value="准备就绪")
        ttk.Label(frame, textvariable=status_var, foreground="blue").pack(pady=(15, 5))
        
        result_var = tk.StringVar()
        result_entry = ttk.Entry(frame, textvariable=result_var, state='readonly', font=("Consolas", 10))
        result_entry.pack(fill=tk.X, pady=5)
        
        def copy_link():
            if result_var.get():
                self.root.clipboard_clear()
                self.root.clipboard_append(result_var.get())
                messagebox.showinfo("复制成功", "链接已复制到剪贴板")
        
        copy_btn = ttk.Button(frame, text="复制链接", command=copy_link, state='disabled')
        copy_btn.pack(pady=5)
        
        def start_upload():
            path = file_path_var.get()
            if not path or not os.path.exists(path):
                messagebox.showerror("错误", "文件不存在")
                return
                
            service = service_var.get()
            status_var.set("正在上传，请稍候...")
            upload_btn.config(state='disabled')
            dialog.update()
            
            def run():
                try:
                    url = ""
                    if "Catbox" in service:
                        url = CloudUploader.upload_to_catbox(path)
                    elif "Litterbox" in service:
                        url = CloudUploader.upload_to_litterbox(path, time='72h')
                    elif "File.io" in service:
                        url = CloudUploader.upload_to_fileio(path)
                        
                    dialog.after(0, lambda: success(url))
                except Exception as e:
                    dialog.after(0, lambda: fail(str(e)))
            
            threading.Thread(target=run, daemon=True).start()
            
        def success(url):
            status_var.set("上传成功!")
            result_var.set(url)
            upload_btn.config(state='normal')
            copy_btn.config(state='normal')
            messagebox.showinfo("成功", f"文件上传成功!\n链接: {url}\n\n该链接可在任何地方访问。")
            
        def fail(msg):
            status_var.set("上传失败")
            result_var.set("")
            upload_btn.config(state='normal')
            messagebox.showerror("上传失败", msg)
            
        upload_btn = ttk.Button(frame, text="开始上传", command=start_upload)
        upload_btn.pack(pady=10)
        
        ttk.Label(frame, text="注意: 请勿上传敏感或隐私文件。", foreground="gray", font=("", 8)).pack(side=tk.BOTTOM, pady=5)

    def generate_toc(self, text):
        """生成目录结构"""
        self.toc_tree.delete(*self.toc_tree.get_children())
        if not text: return
        
        # 常见章节匹配模式
        patterns = [
            r'(^|\n)\s*(第[0-9一二三四五六七八九十百]+[章节回].{0,30})',
            r'(^|\n)\s*(Chapter\s+[0-9IVX]+.{0,30})',
            r'(^|\n)\s*(\d+\.\s+.{0,30})'
        ]
        
        matches = []
        for pat in patterns:
            for m in re.finditer(pat, text):
                matches.append((m.start(), m.group(0).strip()))
        
        # 排序
        matches.sort(key=lambda x: x[0])
        
        # 去重与过滤
        unique_matches = []
        last_pos = -1
        for pos, title in matches:
            if pos > last_pos + 100: # 假设章节间隔至少100字符
                unique_matches.append((pos, title))
                last_pos = pos
                
        # 填充树
        for pos, title in unique_matches:
            self.toc_tree.insert("", "end", text=title, values=(pos,))

    def on_toc_click(self, event):
        """处理目录点击跳转"""
        selected = self.toc_tree.selection()
        if not selected: return
        item = self.toc_tree.item(selected[0])
        pos = int(item['values'][0])
        
        # 跳转原文
        index = f"1.0 + {pos} chars"
        self.original_text.see(index)
        self.original_text.tag_remove("highlight", "1.0", "end")
        self.original_text.tag_add("highlight", index, f"{index} lineend")
        self.original_text.tag_config("highlight", background="yellow")

    def browse_file(self):
        """浏览并选择文件"""
        filename = filedialog.askopenfilename(
            title="选择要翻译的书籍",
            filetypes=self.file_processor.get_file_filter()
        )

        if filename:
            self.file_path_var.set(filename)
            self.load_file_content(filename)

    def load_file_content(self, filepath):
        """加载文件内容"""
        try:
            content = self._read_content_for_path(filepath)
            load_info = self.load_content_into_workspace(
                title=Path(filepath).name,
                content=content,
                filepath=filepath,
                clear_progress_cache=True,
            )

            # 提示信息
            msg = (
                "文件加载成功!\n\n"
                f"字符数: {load_info['char_count']:,}\n"
                f"词数: {load_info['word_count']:,}\n"
                f"预估 Tokens: {load_info['cost_info']['total_estimated_tokens']:,}"
            )
            if load_info['is_large_file']:
                msg += f"\n\n⚠️ 这是一个大文件！\n为了性能，预览窗口仅显示前 {self.preview_limit:,} 字符。\n\n翻译时会使用完整文本。"
            messagebox.showinfo("成功", msg)

        except Exception as e:
            messagebox.showerror("错误", f"加载文件失败:\n{str(e)}")

    # ==================== 批量处理功能 ====================

    def open_batch_window(self):
        """打开批量任务管理窗口"""
        if self.batch_window and self.batch_window.winfo_exists():
            self.batch_window.lift()
            return

        self.batch_window = tk.Toplevel(self.root)
        self.batch_window.title("批量翻译任务")
        self.batch_window.geometry("600x450")
        
        # 顶部说明
        ttk.Label(
            self.batch_window, 
            text="批量添加文件，程序将自动逐个翻译并导出。\n请确保API配额充足或启用了本地模型回退。",
            justify=tk.LEFT, padding=10
        ).pack(fill=tk.X)
        
        # 任务列表
        list_frame = ttk.Frame(self.batch_window, padding=10)
        list_frame.pack(fill=tk.BOTH, expand=True)
        
        self.batch_listbox = tk.Listbox(list_frame, selectmode=tk.EXTENDED)
        self.batch_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        scrollbar = ttk.Scrollbar(list_frame, orient=tk.VERTICAL, command=self.batch_listbox.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.batch_listbox.config(yscrollcommand=scrollbar.set)
        
        # 更新列表显示
        self.update_batch_list()
        
        # 底部按钮
        btn_frame = ttk.Frame(self.batch_window, padding=10)
        btn_frame.pack(fill=tk.X)
        
        ttk.Button(btn_frame, text="添加文件...", command=self.add_batch_files).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="移除选中", command=self.remove_batch_file).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="清空列表", command=lambda: [self.batch_queue.clear(), self.update_batch_list()]).pack(side=tk.LEFT, padx=5)
        
        ttk.Frame(btn_frame).pack(side=tk.LEFT, expand=True) # Spacer
        
        self.batch_start_btn = ttk.Button(btn_frame, text="开始批量处理", command=self.start_batch_processing)
        self.batch_start_btn.pack(side=tk.RIGHT, padx=5)

    def update_batch_list(self):
        """更新批量任务列表显示"""
        if not hasattr(self, 'batch_listbox') or not self.batch_listbox.winfo_exists():
            return
            
        self.batch_listbox.delete(0, tk.END)
        for item in self.batch_queue:
            status_icon = "⏳"
            if item['status'] == 'done': status_icon = "✅"
            elif item['status'] == 'processing': status_icon = "🔄"
            elif item['status'] == 'failed': status_icon = "❌"
            
            self.batch_listbox.insert(tk.END, f"{status_icon} {Path(item['path']).name} ({item['status']})")

    def add_batch_files(self):
        """添加文件到批量队列"""
        filenames = filedialog.askopenfilenames(
            title="选择要批量翻译的文件",
            filetypes=self.file_processor.get_file_filter()
        )
        if filenames:
            for f in filenames:
                # 查重
                if not any(item['path'] == f for item in self.batch_queue):
                    self.batch_queue.append({
                        'path': f,
                        'status': 'pending'
                    })
            self.save_batch_queue()
            self.update_batch_list()

    def remove_batch_file(self):
        """移除选中的批量任务"""
        indices = self.batch_listbox.curselection()
        if not indices: return
        
        # 从后往前删，避免索引偏移
        for i in sorted(indices, reverse=True):
            if i < len(self.batch_queue):
                del self.batch_queue[i]
        self.save_batch_queue()
        self.update_batch_list()

    def start_batch_processing(self):
        """开始批量处理"""
        pending = [item for item in self.batch_queue if item['status'] == 'pending']
        if not pending:
            messagebox.showinfo("提示", "没有待处理的任务")
            return
            
        # 选择导出目录
        self.batch_output_dir = filedialog.askdirectory(title="选择批量导出目录")
        if not self.batch_output_dir:
            return
            
        self.is_batch_mode = True
        self.process_next_batch_file()
        
        if self.batch_window:
            self.batch_window.destroy()
            self.batch_window = None

    def process_next_batch_file(self):
        """处理下一个批量文件"""
        if not self.is_batch_mode:
            return

        # 查找下一个 pending 任务
        next_idx = -1
        for i, item in enumerate(self.batch_queue):
            if item['status'] == 'pending':
                next_idx = i
                break
        
        if next_idx == -1:
            self.is_batch_mode = False
            messagebox.showinfo("批量完成", "所有批量任务已处理完毕！")
            return

        # 标记状态
        self.batch_queue[next_idx]['status'] = 'processing'
        self.save_batch_queue()
        file_path = self.batch_queue[next_idx]['path']
        
        # 加载文件
        self.file_path_var.set(file_path)
        # 自动清理旧状态
        self.clear_all_internal(skip_ui_confirm=True)
        self.load_file_content(file_path)
        
        # 开始翻译
        # 使用 root.after 确保 UI 更新后再开始
        self.root.after(1000, self.start_translation)

    def clear_all_internal(self, skip_ui_confirm=False):
        """内部清空方法，供批量模式调用"""
        if not skip_ui_confirm and not messagebox.askyesno("确认", "确定要清空所有内容吗?"):
            return

        self.file_path_var.set("")
        self.current_text = ""
        self.original_text.delete('1.0', tk.END)
        self.reset_translation_state()
        self.reset_analysis_state()
        self.progress_text_var.set("就绪")
        self.docx_handler = None

    def update_concurrency_label(self, val):
        """更新并发数标签和提示"""
        v = int(float(val))
        if v == 1:
            self.concurrency_label.config(text=f"{v} (高质量)")
            self.concurrency_hint_var.set("单线程：启用上下文记忆，翻译质量最高")
        else:
            self.concurrency_label.config(text=f"{v} (高速)")
            self.concurrency_hint_var.set("多线程：速度快，但无上下文记忆（推荐小说/大文件）")

    def on_api_type_change(self, event=None):
        """API类型改变时更新状态"""
        self.update_api_status()

    def update_api_status(self):
        """更新API配置状态"""
        api_type = self.get_current_api_type()
        if self._provider_ready_for_gui(api_type):
            self.api_status_var.set("已配置并可用")
            self.api_status_label.config(foreground="green")
        else:
            reason = provider_error_message(
                api_type,
                api_configs=self.api_configs,
                custom_local_models=self.custom_local_models,
                support_flags=self._provider_support_flags(),
            )
            self.api_status_var.set(reason or "未配置")
            self.api_status_label.config(foreground="orange")

    def get_current_api_type(self):
        """获取当前选择的API类型"""
        return self.get_translation_api_type()

    def _provider_support_flags(self):
        return {
            "gemini": GEMINI_SUPPORT,
            "openai": OPENAI_SUPPORT,
            "claude": CLAUDE_SUPPORT,
            "requests": REQUESTS_SUPPORT,
        }

    def _provider_ready_for_gui(self, api_type):
        """统一判断 GUI 当前 provider 是否可用。"""
        return provider_ready(
            api_type,
            api_configs=self.api_configs,
            custom_local_models=self.custom_local_models,
            support_flags=self._provider_support_flags(),
        )

    def _ensure_provider_ready_or_prompt(self, api_type):
        """确保 provider 已就绪，否则提示并引导用户修复配置。"""
        if self._provider_ready_for_gui(api_type):
            return True

        reason = provider_error_message(
            api_type,
            api_configs=self.api_configs,
            custom_local_models=self.custom_local_models,
            support_flags=self._provider_support_flags(),
        ) or "当前提供商尚未配置完成"

        if reason.startswith("缺少 "):
            messagebox.showerror("错误", reason)
            return False

        messagebox.showwarning("警告", reason)
        if api_type in self.custom_local_models:
            self.open_edit_local_model_dialog(api_type)
        else:
            self.open_api_config(api_type)
        return False

    def _read_content_for_path(self, filepath):
        """统一通过 FileProcessor 读取文件内容。"""
        def update_progress(msg):
            self.progress_text_var.set(msg)
            self.root.update()

        return self.file_processor.read_file(filepath, progress_callback=update_progress)

    def _init_docx_handler_if_needed(self, filepath):
        """按需初始化 DOCX 处理器。"""
        self.docx_handler = None
        if not filepath or not str(filepath).lower().endswith('.docx'):
            return

        try:
            self.docx_handler = DocxHandler(filepath)
        except Exception as e:
            print(f"DOCX 初始化失败: {e}")
            self.docx_handler = None

    def reset_translation_state(self):
        """重置翻译运行态，不影响当前原文内容。"""
        self.translated_text = ""
        self.source_segments = []
        self.translated_segments = []
        self.failed_segments = []
        self.selected_failed_index = None
        self.resume_from_index = 0
        self.lm_studio_fallback_active = False
        self.consecutive_failures = 0
        self.paused_due_to_failures = False
        self.show_full_text = False

        if hasattr(self, 'translated_text_widget'):
            self.translated_text_widget.delete('1.0', tk.END)
        if hasattr(self, 'progress_var'):
            self.progress_var.set(0)
        if hasattr(self, 'file_info_var'):
            self.file_info_var.set("")
        if hasattr(self, 'toggle_preview_btn'):
            self.toggle_preview_btn.config(state='disabled', text="显示完整原文")
        if hasattr(self, 'refresh_failed_segments_view'):
            self.refresh_failed_segments_view()
        if hasattr(self, 'update_comparison_view'):
            self.update_comparison_view()

    def reset_analysis_state(self):
        """重置解析运行态。"""
        self.analysis_segments = []
        if hasattr(self, 'analysis_text'):
            self.analysis_text.delete('1.0', tk.END)
        if hasattr(self, 'analysis_listbox'):
            self.analysis_listbox.delete(0, tk.END)
        if hasattr(self, 'analysis_status_var'):
            self.analysis_status_var.set("翻译完成后可进行解析")

    def load_content_into_workspace(self, title, content, filepath=None, clear_progress_cache=True):
        """统一将内容载入工作区。"""
        if not content:
            raise ValueError("文件内容为空")

        self.reset_translation_state()
        self.reset_analysis_state()

        self.file_path_var.set(filepath or title or "")
        self.current_text = content
        self.text_signature = self.compute_text_signature(content)
        self.generate_toc(content)
        self.update_text_display()

        char_count = len(content)
        word_count = len(content.split())
        is_large_file = char_count > self.preview_limit
        display_name = Path(filepath).name if filepath else (title or "未命名内容")

        if is_large_file:
            self.file_info_var.set(
                f"⚠️ {display_name} ({char_count:,} 字符) - 仅显示前 {self.preview_limit:,} 字符"
            )
            self.toggle_preview_btn.config(state='normal')
        else:
            self.file_info_var.set(f"✓ 已加载 {display_name} ({char_count:,} 字符)")
            self.toggle_preview_btn.config(state='disabled')

        self.progress_text_var.set(f"已载入内容 | 字符数: {char_count:,} | 词数: {word_count:,}")

        model_name = self.api_configs.get(self.get_translation_api_type(), {}).get('model', 'unknown')
        cost_info = CostEstimator.calculate_cost(model_name, content)
        self.cost_var.set(
            f"预估成本: ${cost_info['cost_usd']} (Tokens: {cost_info['total_estimated_tokens']:,})"
        )

        self._init_docx_handler_if_needed(filepath)
        if self.docx_handler:
            self.file_info_var.set(self.file_info_var.get() + " [DOCX 格式保留已就绪]")

        if clear_progress_cache:
            self.clear_progress_cache()

        return {
            "char_count": char_count,
            "word_count": word_count,
            "is_large_file": is_large_file,
            "cost_info": cost_info,
            "display_name": display_name,
        }

    def get_target_language(self):
        """获取用户设置的目标语言"""
        target = (self.target_language_var.get() or "").strip()
        return target if target else DEFAULT_TARGET_LANGUAGE

    def is_target_language_chinese(self, target_language=None):
        """判断目标语言是否为中文"""
        target = (target_language or self.get_target_language() or "").lower()
        return any(key in target for key in ["中文", "汉语", "chinese", "zh"])

    def is_target_language_english(self, target_language=None):
        """判断目标语言是否为英文"""
        target = (target_language or self.get_target_language() or "").lower()
        return any(key in target for key in ["英文", "英语", "english", "en"])

    def compute_text_signature(self, text):
        """计算文本签名用于断点恢复"""
        return hashlib.md5(text.encode('utf-8')).hexdigest() if text else None

    def save_progress_cache(self):
        """保存当前翻译进度到磁盘"""
        try:
            if not self.current_text or not self.source_segments:
                return

            data = {
                'file_path': self.file_path_var.get(),
                'signature': self.text_signature,
                'source_segments': self.source_segments,
                'translated_segments': self.translated_segments,
                'failed_segments': self.failed_segments,
                'lm_studio_fallback_active': self.lm_studio_fallback_active,
                'target_language': self.get_target_language(),
                'resume_from_index': len(self.translated_segments)
            }
            self.progress_cache_path.parent.mkdir(parents=True, exist_ok=True)
            with open(self.progress_cache_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False)
        except Exception as e:
            print(f"保存进度缓存失败: {e}")

    def clear_progress_cache(self):
        """清除翻译进度缓存"""
        try:
            if self.progress_cache_path.exists():
                self.progress_cache_path.unlink()
        except Exception as e:
            print(f"清除进度缓存失败: {e}")

    def try_resume_cached_progress(self):
        """启动时检查并询问是否恢复未完成进度"""
        if not self.progress_cache_path.exists():
            return

        try:
            with open(self.progress_cache_path, 'r', encoding='utf-8') as f:
                cache = json.load(f)
        except Exception as e:
            print(f"读取进度缓存失败: {e}")
            return

        file_path = cache.get('file_path')
        signature = cache.get('signature')
        if not file_path or not Path(file_path).exists():
            print("进度缓存对应的文件不存在，已忽略")
            self.clear_progress_cache()
            return

        try:
            content = self._read_content_for_path(file_path)
        except Exception as e:
            print(f"读取缓存文件失败: {e}")
            self.clear_progress_cache()
            return

        current_sig = self.compute_text_signature(content)
        if signature != current_sig:
            print("文件内容已变化，无法恢复进度，已清除缓存")
            self.clear_progress_cache()
            return

        self.load_content_into_workspace(
            title=Path(file_path).name,
            content=content,
            filepath=file_path,
            clear_progress_cache=False,
        )

        # 恢复状态
        self.text_signature = signature
        self.source_segments = cache.get('source_segments', [])
        self.translated_segments = cache.get('translated_segments', [])
        self.failed_segments = cache.get('failed_segments', [])
        self.lm_studio_fallback_active = cache.get('lm_studio_fallback_active', False)
        self.resume_from_index = cache.get('resume_from_index', len(self.translated_segments))
        cached_target = cache.get('target_language')
        if cached_target:
            self.target_language_var.set(cached_target)

        # 更新界面显示
        self.translated_text = "\n\n".join(self.translated_segments)
        self.update_translated_text(self.translated_text)
        self.refresh_failed_segments_view()

        total_segments = len(self.source_segments) or 1
        progress = (len(self.translated_segments) / total_segments) * 100
        self.progress_var.set(progress)
        self.progress_text_var.set(f"检测到未完成的翻译进度（{len(self.translated_segments)}/{total_segments} 段）")

        continue_resume = messagebox.askyesno(
            "继续未完成的翻译",
            f"检测到未完成的翻译任务：\n文件: {Path(file_path).name}\n进度: {len(self.translated_segments)}/{total_segments}\n\n是否继续？"
        )
        if not continue_resume:
            # 放弃恢复，清理缓存并重置状态
            self.clear_progress_cache()
            self.reset_translation_state()
            self.progress_text_var.set("就绪")

    def open_api_config_for(self, purpose='translation'):
        """打开API配置对话框（根据翻译、解析或重试选择）"""
        if purpose == 'translation':
            api_type = self.get_translation_api_type()
        elif purpose == 'analysis':
            api_type = self.get_analysis_api_type()
        else:
            api_type = self.get_retry_api_type()

        # 如果是自定义本地模型，打开编辑对话框
        if api_type in self.custom_local_models:
            self.open_edit_local_model_dialog(api_type)
        else:
            self.open_api_config(api_type)

    def open_add_local_model_dialog(self):
        """打开添加本地模型对话框"""
        dialog = tk.Toplevel(self.root)
        dialog.title("添加本地模型")
        dialog.geometry("480x320")
        dialog.transient(self.root)
        dialog.grab_set()

        frame = ttk.Frame(dialog, padding="20")
        frame.pack(fill=tk.BOTH, expand=True)
        frame.columnconfigure(1, weight=1)

        # 模型显示名称
        ttk.Label(frame, text="显示名称:").grid(row=0, column=0, sticky=tk.W, pady=5)
        name_var = tk.StringVar()
        ttk.Entry(frame, textvariable=name_var, width=40).grid(row=0, column=1, sticky=(tk.W, tk.E), pady=5, padx=5)

        # Base URL
        ttk.Label(frame, text="Base URL:").grid(row=1, column=0, sticky=tk.W, pady=5)
        url_var = tk.StringVar(value="http://127.0.0.1:1234/v1")
        ttk.Entry(frame, textvariable=url_var, width=40).grid(row=1, column=1, sticky=(tk.W, tk.E), pady=5, padx=5)

        # Model ID
        ttk.Label(frame, text="Model ID:").grid(row=2, column=0, sticky=tk.W, pady=5)
        model_var = tk.StringVar()
        ttk.Entry(frame, textvariable=model_var, width=40).grid(row=2, column=1, sticky=(tk.W, tk.E), pady=5, padx=5)

        # API Key (可选)
        ttk.Label(frame, text="API Key:").grid(row=3, column=0, sticky=tk.W, pady=5)
        key_var = tk.StringVar(value="lm-studio")
        ttk.Entry(frame, textvariable=key_var, width=40).grid(row=3, column=1, sticky=(tk.W, tk.E), pady=5, padx=5)

        # 帮助文本
        ttk.Label(
            frame,
            text="提示: 本地模型使用OpenAI兼容接口格式\n例如 LM Studio, Ollama, vLLM 等",
            foreground="gray",
            justify=tk.LEFT
        ).grid(row=4, column=0, columnspan=2, sticky=tk.W, pady=10)

        def test_connection():
            """测试本地模型连接"""
            if not OPENAI_SUPPORT:
                messagebox.showerror("错误", "缺少 openai 库，无法测试连接")
                return

            test_url = url_var.get().strip()
            test_model = model_var.get().strip()
            test_key = key_var.get().strip() or "lm-studio"

            if not test_url or not test_model:
                messagebox.showwarning("警告", "请填写 Base URL 和 Model ID")
                return

            try:
                client = openai.OpenAI(api_key=test_key, base_url=test_url)
                response = client.chat.completions.create(
                    model=test_model,
                    messages=[{"role": "user", "content": "Hi"}],
                    max_tokens=5
                )
                messagebox.showinfo("成功", f"连接测试成功!\n响应: {response.choices[0].message.content[:50]}")
            except Exception as e:
                messagebox.showerror("连接失败", f"无法连接到本地模型:\n{str(e)}")

        def save_model():
            name = name_var.get().strip()
            url = url_var.get().strip()
            model = model_var.get().strip()
            key = key_var.get().strip() or "lm-studio"

            if not name:
                messagebox.showwarning("警告", "请输入显示名称")
                return
            if not url:
                messagebox.showwarning("警告", "请输入 Base URL")
                return
            if not model:
                messagebox.showwarning("警告", "请输入 Model ID")
                return

            try:
                self.add_custom_local_model(
                    name=name,
                    display_name=name,
                    base_url=url,
                    model_id=model,
                    api_key=key
                )
                messagebox.showinfo("成功", f"本地模型 '{name}' 已添加")
                dialog.destroy()
            except Exception as e:
                messagebox.showerror("错误", str(e))

        btn_frame = ttk.Frame(frame)
        btn_frame.grid(row=5, column=0, columnspan=2, pady=20)
        ttk.Button(btn_frame, text="测试连接", command=test_connection).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="保存", command=save_model).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="取消", command=dialog.destroy).pack(side=tk.LEFT, padx=5)

    def open_manage_local_models_dialog(self):
        """打开管理本地模型对话框"""
        dialog = tk.Toplevel(self.root)
        dialog.title("管理本地模型")
        dialog.geometry("550x400")
        dialog.transient(self.root)
        dialog.grab_set()

        frame = ttk.Frame(dialog, padding="15")
        frame.pack(fill=tk.BOTH, expand=True)
        frame.columnconfigure(0, weight=1)
        frame.rowconfigure(1, weight=1)

        ttk.Label(frame, text="已添加的本地模型:", font=('', 10, 'bold')).grid(row=0, column=0, sticky=tk.W, pady=(0, 5))

        # 模型列表框架
        list_frame = ttk.Frame(frame)
        list_frame.grid(row=1, column=0, sticky=(tk.N, tk.S, tk.E, tk.W), pady=(0, 10))
        list_frame.columnconfigure(0, weight=1)
        list_frame.rowconfigure(0, weight=1)

        listbox = tk.Listbox(list_frame, height=12, font=('', 10))
        listbox.grid(row=0, column=0, sticky=(tk.N, tk.S, tk.E, tk.W))

        scrollbar = ttk.Scrollbar(list_frame, orient=tk.VERTICAL, command=listbox.yview)
        scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S))
        listbox.config(yscrollcommand=scrollbar.set)

        # 详情显示
        detail_var = tk.StringVar(value="选择一个模型查看详情")
        detail_label = ttk.Label(frame, textvariable=detail_var, foreground="gray", wraplength=500, justify=tk.LEFT)
        detail_label.grid(row=2, column=0, sticky=tk.W, pady=(0, 10))

        def refresh_list():
            listbox.delete(0, tk.END)
            if not self.custom_local_models:
                listbox.insert(tk.END, "(暂无自定义本地模型)")
                detail_var.set("点击「+ 添加本地模型」按钮添加新模型")
                return

            for key, config in self.custom_local_models.items():
                listbox.insert(tk.END, f"{config['display_name']} ({config['base_url']})")

        def on_select(event=None):
            selection = listbox.curselection()
            if not selection or not self.custom_local_models:
                return

            keys = list(self.custom_local_models.keys())
            if selection[0] >= len(keys):
                return

            key = keys[selection[0]]
            config = self.custom_local_models[key]
            detail_var.set(
                f"名称: {config['display_name']}\n"
                f"Base URL: {config['base_url']}\n"
                f"Model ID: {config['model_id']}\n"
                f"API Key: {config.get('api_key', 'lm-studio')}"
            )

        listbox.bind('<<ListboxSelect>>', on_select)

        def delete_selected():
            selection = listbox.curselection()
            if not selection or not self.custom_local_models:
                messagebox.showinfo("提示", "请先选择要删除的模型")
                return

            keys = list(self.custom_local_models.keys())
            if selection[0] >= len(keys):
                return

            key = keys[selection[0]]
            config = self.custom_local_models[key]

            if messagebox.askyesno("确认删除", f"确定删除模型 '{config['display_name']}'?"):
                self.remove_custom_local_model(key)
                refresh_list()
                detail_var.set("选择一个模型查看详情")

        def edit_selected():
            selection = listbox.curselection()
            if not selection or not self.custom_local_models:
                messagebox.showinfo("提示", "请先选择要编辑的模型")
                return

            keys = list(self.custom_local_models.keys())
            if selection[0] >= len(keys):
                return

            key = keys[selection[0]]
            dialog.destroy()
            self.open_edit_local_model_dialog(key)

        refresh_list()

        btn_frame = ttk.Frame(frame)
        btn_frame.grid(row=3, column=0, sticky=tk.W)
        ttk.Button(btn_frame, text="编辑", command=edit_selected).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="删除", command=delete_selected).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="关闭", command=dialog.destroy).pack(side=tk.LEFT, padx=5)

    def open_edit_local_model_dialog(self, model_key):
        """打开编辑本地模型对话框"""
        if model_key not in self.custom_local_models:
            messagebox.showerror("错误", f"模型 '{model_key}' 不存在")
            return

        config = self.custom_local_models[model_key]

        dialog = tk.Toplevel(self.root)
        dialog.title(f"编辑本地模型: {config['display_name']}")
        dialog.geometry("480x320")
        dialog.transient(self.root)
        dialog.grab_set()

        frame = ttk.Frame(dialog, padding="20")
        frame.pack(fill=tk.BOTH, expand=True)
        frame.columnconfigure(1, weight=1)

        # 模型显示名称
        ttk.Label(frame, text="显示名称:").grid(row=0, column=0, sticky=tk.W, pady=5)
        name_var = tk.StringVar(value=config['display_name'])
        ttk.Entry(frame, textvariable=name_var, width=40).grid(row=0, column=1, sticky=(tk.W, tk.E), pady=5, padx=5)

        # Base URL
        ttk.Label(frame, text="Base URL:").grid(row=1, column=0, sticky=tk.W, pady=5)
        url_var = tk.StringVar(value=config['base_url'])
        ttk.Entry(frame, textvariable=url_var, width=40).grid(row=1, column=1, sticky=(tk.W, tk.E), pady=5, padx=5)

        # Model ID
        ttk.Label(frame, text="Model ID:").grid(row=2, column=0, sticky=tk.W, pady=5)
        model_var = tk.StringVar(value=config['model_id'])
        ttk.Entry(frame, textvariable=model_var, width=40).grid(row=2, column=1, sticky=(tk.W, tk.E), pady=5, padx=5)

        # API Key
        ttk.Label(frame, text="API Key:").grid(row=3, column=0, sticky=tk.W, pady=5)
        key_var = tk.StringVar(value=config.get('api_key', 'lm-studio'))
        ttk.Entry(frame, textvariable=key_var, width=40).grid(row=3, column=1, sticky=(tk.W, tk.E), pady=5, padx=5)

        def test_connection():
            if not OPENAI_SUPPORT:
                messagebox.showerror("错误", "缺少 openai 库")
                return
            try:
                client = openai.OpenAI(
                    api_key=key_var.get().strip() or "lm-studio",
                    base_url=url_var.get().strip()
                )
                response = client.chat.completions.create(
                    model=model_var.get().strip(),
                    messages=[{"role": "user", "content": "Hi"}],
                    max_tokens=5
                )
                messagebox.showinfo("成功", "连接测试成功!")
            except Exception as e:
                messagebox.showerror("连接失败", str(e))

        def save_changes():
            self.custom_local_models[model_key] = {
                'display_name': name_var.get().strip(),
                'base_url': url_var.get().strip(),
                'model_id': model_var.get().strip(),
                'api_key': key_var.get().strip() or "lm-studio",
                'created_at': config.get('created_at', '')
            }
            self.refresh_api_dropdowns()
            self.save_config()
            messagebox.showinfo("成功", "模型配置已更新")
            dialog.destroy()

        btn_frame = ttk.Frame(frame)
        btn_frame.grid(row=5, column=0, columnspan=2, pady=20)
        ttk.Button(btn_frame, text="测试连接", command=test_connection).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="保存", command=save_changes).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="取消", command=dialog.destroy).pack(side=tk.LEFT, padx=5)

    def open_api_config(self, api_type=None):
        """打开API配置对话框"""
        if api_type is None:
            api_type = self.get_current_api_type()
        config = self.api_configs[api_type]

        title_map = {
            'gemini': "Gemini API",
            'openai': "OpenAI API",
            'claude': "Claude API",
            'deepseek': "DeepSeek API",
            'lm_studio': "本地 LM Studio",
            'custom': "自定义 API",
        }

        config_window = tk.Toplevel(self.root)
        config_window.title(f"{title_map.get(api_type, api_type)} 配置")
        config_window.geometry("500x350")
        config_window.transient(self.root)
        config_window.grab_set()

        frame = ttk.Frame(config_window, padding="20")
        frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        frame.columnconfigure(1, weight=1)

        # API Key
        ttk.Label(frame, text="API Key:").grid(row=0, column=0, sticky=tk.W, pady=5)
        api_key_var = tk.StringVar(value=config.get('api_key', ''))
        ttk.Entry(frame, textvariable=api_key_var, width=50).grid(
            row=0, column=1, sticky=(tk.W, tk.E), pady=5, padx=5
        )

        # Model
        ttk.Label(frame, text="模型:").grid(row=1, column=0, sticky=tk.W, pady=5)
        model_var = tk.StringVar(value=config.get('model', ''))
        ttk.Entry(frame, textvariable=model_var, width=50).grid(
            row=1, column=1, sticky=(tk.W, tk.E), pady=5, padx=5
        )

        # Base URL (for OpenAI compatible APIs)
        if api_type in ['openai', 'custom', 'lm_studio', 'deepseek']:
            ttk.Label(frame, text="Base URL:").grid(row=2, column=0, sticky=tk.W, pady=5)
            base_url_var = tk.StringVar(value=config.get('base_url', ''))
            ttk.Entry(frame, textvariable=base_url_var, width=50).grid(
                row=2, column=1, sticky=(tk.W, tk.E), pady=5, padx=5
            )
            ttk.Label(
                frame,
                text="(可选，用于自定义服务或代理)",
                foreground="gray"
            ).grid(row=3, column=1, sticky=tk.W, pady=5)
        else:
            base_url_var = tk.StringVar(value='')

        # 说明文本
        help_text = {
            'gemini': "请在 Google AI Studio 获取 API Key\n模型示例: gemini-2.5-flash, gemini-2.5-pro",
            'openai': "请在 OpenAI 获取 API Key\n模型示例: gpt-3.5-turbo, gpt-4",
            'claude': "请在 Anthropic Console 获取 API Key\n模型示例: claude-3-haiku-20240307",
            'deepseek': "请在 DeepSeek 开放平台获取 API Key\n模型示例: deepseek-chat",
            'custom': (
                "输入兼容OpenAI API格式的自定义服务\n"
                "Base URL示例: https://api.example.com/v1\n"
                "本地LM Studio示例: http://127.0.0.1:1234/v1 (模型如 qwen2.5-7b-instruct-1m)"
            ),
            'lm_studio': (
                "连接本地 LM Studio 提供的 OpenAI 兼容接口\n"
                "默认地址: http://127.0.0.1:1234/v1\n"
                "API Key 可留空，程序会自动使用占位值\n"
                "请确保 LM Studio Server 已启动并加载目标模型"
            )
        }

        help_label = ttk.Label(
            frame,
            text=help_text.get(api_type, ''),
            foreground="gray",
            justify=tk.LEFT
        )
        help_label.grid(row=4, column=0, columnspan=2, sticky=tk.W, pady=10)

        # 测试连接按钮
        def test_connection():
            """测试API连接"""
            test_config = {
                'api_key': api_key_var.get().strip(),
                'model': model_var.get().strip(),
                'base_url': base_url_var.get().strip(),
            }
            ready, reason = validate_builtin_provider(
                api_type,
                test_config,
                support_flags=self._provider_support_flags(),
            )
            if not ready:
                messagebox.showwarning("警告", reason)
                return

            # 显示测试中提示
            test_btn.config(state='disabled', text="测试中...")
            config_window.update()

            try:
                # 使用 TranslationEngine 的测试功能
                test_engine = TranslationEngine()

                test_engine.add_api_config(api_type, APIConfig(
                    provider=self._builtin_provider_enum(api_type),
                    api_key=test_config['api_key'],
                    model=test_config['model'],
                    base_url=test_config['base_url'],
                    temperature=0.2
                ))
                
                success, msg = test_engine.test_connection(api_type)
                
                if success:
                    messagebox.showinfo("成功", f"✓ API连接测试成功！\n\n{msg}")
                else:
                    messagebox.showerror("测试失败", f"✗ API连接测试失败\n\n{msg}")

            except Exception as e:
                messagebox.showerror("测试失败", f"✗ API连接测试失败\n\n错误: {str(e)}\n\n请检查配置是否正确")

            finally:
                test_btn.config(state='normal', text="测试连接")

        # 保存按钮
        def save_config():
            new_config = {
                'api_key': api_key_var.get().strip(),
                'model': model_var.get().strip(),
                'base_url': base_url_var.get().strip(),
            }
            ready, reason = validate_builtin_provider(
                api_type,
                new_config,
                support_flags=self._provider_support_flags(),
            )
            if not ready:
                messagebox.showwarning("警告", reason)
                return

            # 保存配置
            self.api_configs[api_type]['api_key'] = new_config['api_key']
            self.api_configs[api_type]['model'] = new_config['model']
            if api_type in ['openai', 'custom', 'lm_studio', 'deepseek']:
                self.api_configs[api_type]['base_url'] = new_config['base_url']

            # 自动保存到文件
            self.save_config(show_message=True)
            self.update_api_status()
            config_window.destroy()
            messagebox.showinfo("成功", "✓ API配置已保存\n✓ 已自动创建备份\n\n配置将在下次启动时自动加载")

        button_frame = ttk.Frame(frame)
        button_frame.grid(row=5, column=0, columnspan=2, pady=20)

        test_btn = ttk.Button(button_frame, text="测试连接", command=test_connection)
        test_btn.grid(row=0, column=0, padx=5)
        ttk.Button(button_frame, text="保存", command=save_config).grid(row=0, column=1, padx=5)
        ttk.Button(button_frame, text="取消", command=config_window.destroy).grid(row=0, column=2, padx=5)

    def merge_api_configs(self, incoming_configs):
        """将磁盘配置与默认值合并，确保新字段有默认值"""
        incoming_configs = incoming_configs or {}

        for name, defaults in DEFAULT_API_CONFIGS.items():
            merged = deepcopy(defaults)
            merged.update(incoming_configs.get(name, {}))
            self.api_configs[name] = merged

        # 保留未知的扩展配置，避免意外丢失
        for extra_key, extra_val in incoming_configs.items():
            if extra_key not in self.api_configs:
                self.api_configs[extra_key] = extra_val

    def migrate_config_v1_to_v2(self, old_config):
        """将v1配置迁移到v2格式"""
        new_config = {
            'version': CONFIG_VERSION,
            'api_configs': old_config.get('api_configs', {}),
            'custom_local_models': {},
            'target_language': old_config.get('target_language', DEFAULT_TARGET_LANGUAGE),
            'selected_translation_api': 'Gemini API',
            'selected_analysis_api': 'Gemini API'
        }
        return new_config

    def get_all_available_apis(self):
        """获取所有可用的API列表（内置+自定义本地模型）"""
        apis = []

        # 内置API
        if GEMINI_SUPPORT:
            apis.append(("Gemini API", "gemini", "builtin"))
        if OPENAI_SUPPORT:
            apis.append(("OpenAI API", "openai", "builtin"))
            apis.append(("DeepSeek API", "deepseek", "builtin"))
            apis.append(("本地 LM Studio", "lm_studio", "builtin"))
        if CLAUDE_SUPPORT:
            apis.append(("Claude API", "claude", "builtin"))
        if REQUESTS_SUPPORT:
            apis.append(("自定义API", "custom", "builtin"))

        # 自定义本地模型
        for key, config in self.custom_local_models.items():
            display_name = config.get('display_name', key)
            apis.append((f"[本地] {display_name}", key, "custom_local"))

        return apis

    def add_custom_local_model(self, name, display_name, base_url, model_id, api_key="lm-studio"):
        """添加自定义本地模型"""
        from datetime import datetime

        # 生成唯一键名
        key = name.lower().replace(" ", "_").replace("-", "_")
        if key in self.api_configs or key in self.custom_local_models:
            raise ValueError(f"模型名称 '{name}' 已存在")

        self.custom_local_models[key] = {
            'display_name': display_name,
            'base_url': base_url,
            'model_id': model_id,
            'api_key': api_key,
            'created_at': datetime.now().isoformat()
        }

        # 刷新下拉框
        self.refresh_api_dropdowns()
        self.save_config()
        return key

    def remove_custom_local_model(self, key):
        """删除自定义本地模型"""
        if key in self.custom_local_models:
            del self.custom_local_models[key]
            self.refresh_api_dropdowns()
            self.save_config()
            return True
        return False

    def refresh_api_dropdowns(self):
        """刷新翻译和解析的API下拉框"""
        apis = self.get_all_available_apis()
        api_names = [name for name, _, _ in apis]

        if hasattr(self, 'translation_api_combo'):
            current_trans = self.translation_api_var.get()
            self.translation_api_combo['values'] = api_names
            # 保持当前选择，如果仍然有效
            if current_trans not in api_names and api_names:
                self.translation_api_var.set(api_names[0])

        if hasattr(self, 'analysis_api_combo'):
            current_analysis = self.analysis_api_var.get()
            self.analysis_api_combo['values'] = api_names
            if current_analysis not in api_names and api_names:
                self.analysis_api_var.set(api_names[0])

        if hasattr(self, 'retry_api_combo'):
            current_retry = self.retry_api_var.get()
            self.retry_api_combo['values'] = api_names
            if current_retry not in api_names and api_names:
                preferred_retry = "本地 LM Studio" if "本地 LM Studio" in api_names else self.translation_api_var.get()
                if preferred_retry not in api_names:
                    preferred_retry = api_names[0]
                self.retry_api_var.set(preferred_retry)

    def _map_api_name_to_key(self, api_name):
        """将显示名称映射到API键"""
        if not api_name:
            return "gemini"

        # 检查是否为自定义本地模型
        if api_name.startswith("[本地] "):
            display_name = api_name[5:]  # 去掉"[本地] "前缀
            for key, config in self.custom_local_models.items():
                if config.get('display_name') == display_name:
                    return key

        # 内置API映射
        api_map = {
            "Gemini API": "gemini",
            "OpenAI API": "openai",
            "Claude API": "claude",
            "DeepSeek API": "deepseek",
            "本地 LM Studio": "lm_studio",
            "自定义API": "custom"
        }
        return api_map.get(api_name, "gemini")

    def get_translation_api_type(self):
        """获取当前选择的翻译API类型"""
        api_name = self.translation_api_var.get()
        return self._map_api_name_to_key(api_name)

    def get_analysis_api_type(self):
        """获取当前选择的解析API类型"""
        api_name = self.analysis_api_var.get()
        return self._map_api_name_to_key(api_name)

    def get_retry_api_type(self):
        """获取失败段落重试时选择的API类型"""
        api_name = self.retry_api_var.get()
        return self._map_api_name_to_key(api_name)

    def save_config(self, show_message=False):
        """保存配置到用户配置目录。"""
        try:
            self.config_manager.set('api_configs', deepcopy(self.api_configs), save=False)
            self.config_manager.set('custom_local_models', deepcopy(self.custom_local_models), save=False)
            self.config_manager.set('target_language', self.get_target_language(), save=False)
            self.config_manager.set('selected_translation_api', self.translation_api_var.get(), save=False)
            self.config_manager.set('selected_analysis_api', self.analysis_api_var.get(), save=False)
            self.config_manager.set('selected_retry_api', self.retry_api_var.get(), save=False)
            self.config_manager.save(create_backup=True)

            if show_message:
                print(f"✓ 配置已自动保存: {self.config_manager.config_path}")
        except Exception as e:
            error_msg = f"保存配置失败: {e}"
            print(error_msg)
            if show_message:
                messagebox.showerror("错误", error_msg)

    def backup_config(self):
        """兼容旧调用，备份已由 ConfigManager.save 统一处理。"""
        try:
            self.config_manager.save(create_backup=True)
        except Exception as e:
            print(f"备份配置失败: {e}")

    def load_config(self):
        """从用户配置目录加载配置。"""
        try:
            loaded_config = self.config_manager.get_all()
            api_config_section = loaded_config.get('api_configs', {})
            self.merge_api_configs(api_config_section)

            self.custom_local_models = loaded_config.get('custom_local_models', {})
            target_language = loaded_config.get('target_language', DEFAULT_TARGET_LANGUAGE)
            self.target_language_var.set(target_language or DEFAULT_TARGET_LANGUAGE)

            saved_trans_api = loaded_config.get('selected_translation_api', 'Gemini API')
            saved_analysis_api = loaded_config.get('selected_analysis_api', 'Gemini API')
            saved_retry_api = loaded_config.get('selected_retry_api')
            self.translation_api_var.set(saved_trans_api)
            self.analysis_api_var.set(saved_analysis_api)
            default_retry_api = "本地 LM Studio" if OPENAI_SUPPORT else saved_trans_api
            self.retry_api_var.set(saved_retry_api or default_retry_api)

            self.update_api_status()
            self.refresh_api_dropdowns()
            ready_builtin = [
                name for name in self.api_configs
                if self._provider_ready_for_gui(name)
            ]
            ready_local = [
                name for name in self.custom_local_models
                if self._provider_ready_for_gui(name)
            ]
            print(
                f"✓ 配置已加载: {len(ready_builtin)} 个内置API可用, "
                f"{len(ready_local)} 个自定义本地模型可用"
            )
        except Exception as e:
            error_msg = f"加载配置失败: {e}"
            print(error_msg)
            if self.restore_from_backup():
                print("✓ 已从备份恢复配置")
            else:
                messagebox.showwarning("警告", f"{error_msg}\n将使用默认配置")

    def restore_from_backup(self):
        """从最新备份恢复配置。"""
        try:
            if not self.config_manager._restore_from_backup():
                return False
            self.load_config()
            return True
        except Exception as e:
            print(f"从备份恢复失败: {e}")
            return False

    def on_closing(self):
        """程序退出时的处理（自动保存配置）"""
        # 如果正在翻译，询问用户
        if self.is_translating:
            if not messagebox.askyesno("确认退出", "翻译正在进行中，确定要退出吗？\n\n配置将自动保存"):
                return

        # 自动保存配置
        try:
            self.save_config(show_message=False)
            print("✓ 配置已自动保存")
        except Exception as e:
            print(f"保存配置时出错: {e}")

        # 关闭窗口
        self.root.destroy()

    def start_translation(self):
        """开始翻译"""
        if not self.current_text:
            messagebox.showwarning("警告", "请先加载要翻译的文件")
            return

        api_type = self.get_translation_api_type()
        if not self._ensure_provider_ready_or_prompt(api_type):
            return

        # 计算签名用于断点恢复判断
        current_signature = self.compute_text_signature(self.current_text)
        resume_possible = (
            self.text_signature == current_signature
            and self.source_segments
            and 0 < len(self.translated_segments) < len(self.source_segments)
        )

        # 是否从断点继续
        self.resume_from_index = 0
        if resume_possible:
            resume = messagebox.askyesno(
                "继续翻译",
                f"检测到上次未完成的翻译，是否从第 {len(self.translated_segments) + 1} 段继续？"
            )
            if resume:
                self.resume_from_index = len(self.translated_segments)
                # 确保译文长度与起始段对齐
                if len(self.translated_segments) > self.resume_from_index:
                    self.translated_segments = self.translated_segments[:self.resume_from_index]
            else:
                self.translated_segments = []
                self.source_segments = []
                self.failed_segments = []
        else:
            self.translated_segments = []
            self.source_segments = []
            self.failed_segments = []

        # 开始翻译
        self.lm_studio_fallback_active = False
        self.consecutive_failures = 0
        self.paused_due_to_failures = False
        self.is_translating = True
        self.translate_btn.config(state='disabled')
        self.stop_btn.config(state='normal')
        self.progress_var.set(
            (self.resume_from_index / max(len(self.source_segments), 1)) * 100
            if self.resume_from_index and self.source_segments else 0
        )
        if not self.resume_from_index:
            self.translated_text = ""
            self.translated_text_widget.delete('1.0', tk.END)
        self.failed_segments = []
        self.selected_failed_index = None
        self.refresh_failed_segments_view()

        # 在新线程中执行翻译
        self.translation_thread = threading.Thread(target=self.translate_text, daemon=True)
        self.translation_thread.start()

    def stop_translation(self):
        """停止翻译"""
        self.is_translating = False
        self.translate_btn.config(state='normal')
        self.stop_btn.config(state='disabled')
        self.progress_text_var.set("翻译已停止")

    def sync_engine_config(self):
        """同步配置到翻译引擎"""
        # 清除旧配置
        self.translation_engine.api_configs.clear()
        self.translation_engine.custom_local_models.clear()

        support_flags = self._provider_support_flags()

        # 同步内置 API 配置
        for name, cfg in self.api_configs.items():
            provider_enum = provider_enum_for_name(name)
            if provider_enum is None:
                continue

            if not provider_ready(
                name,
                api_configs=self.api_configs,
                support_flags=support_flags,
            ):
                continue

            self.translation_engine.add_api_config(name, APIConfig(
                provider=provider_enum,
                api_key=cfg.get('api_key', ''),
                model=cfg.get('model', ''),
                base_url=cfg.get('base_url', ''),
                temperature=cfg.get('temperature', 0.2)
            ))

        # 同步自定义本地模型
        for name, cfg in self.custom_local_models.items():
            if not provider_ready(
                name,
                custom_local_models=self.custom_local_models,
                support_flags=support_flags,
            ):
                continue

            self.translation_engine.add_custom_local_model(
                name=name,
                display_name=cfg.get('display_name', name),
                base_url=cfg.get('base_url', ''),
                model_id=cfg.get('model_id', ''),
                api_key=cfg.get('api_key') or 'lm-studio'
            )

        # 设置回退逻辑
        self.translation_engine.fallback_provider = choose_fallback_provider(
            api_configs=self.api_configs,
            custom_local_models=self.custom_local_models,
            support_flags=support_flags,
        )

    def translate_text(self):
        """执行翻译（在后台线程中，支持并发）"""
        try:
            # 同步配置到引擎
            self.sync_engine_config()
            
            # 获取当前翻译API类型
            api_type = self.get_translation_api_type()
            self.consecutive_failures = 0

            # 准备
            self.root.after(0, self.progress_text_var.set, "正在进行文本分段...")

            # 使用 FileProcessor 进行分段
            self.source_segments = self.file_processor.split_text_into_segments(self.current_text, max_length=800)
            total_segments = len(self.source_segments)
            self.text_signature = self.compute_text_signature(self.current_text)
            start_index = min(self.resume_from_index or 0, total_segments)

            self.root.after(0, self.progress_text_var.set, f"文本已分为 {total_segments} 段，准备开始翻译...")
            if start_index:
                self.root.after(
                    0,
                    self.progress_var.set,
                    (start_index / total_segments) * 100 if total_segments else 0
                )
                self.root.after(
                    0,
                    self.progress_text_var.set,
                    f"继续翻译：从第 {start_index + 1} 段开始..."
                )

            # 预填充翻译列表，确保索引对齐
            if len(self.translated_segments) < total_segments:
                self.translated_segments.extend([""] * (total_segments - len(self.translated_segments)))

            # 获取并发设置
            max_workers = self.concurrency_var.get()
            remaining_segments = max(total_segments - start_index, 0)
            max_workers = max(1, min(max_workers, remaining_segments or 1))
            use_context = max_workers == 1  # 只有单线程模式才启用上下文
            
            # 定义单个任务函数
            def process_segment(idx):
                if not self.is_translating or self.paused_due_to_failures:
                    return None
                    
                segment = self.source_segments[idx]
                
                # 获取上下文（仅单线程有效）
                context = None
                if use_context and idx > 0:
                    prev_trans = self.translated_segments[idx-1]
                    # 确保前一段已翻译且不是错误信息
                    if prev_trans and not prev_trans.startswith("["):
                        context = prev_trans

                try:
                    result = self.translate_segment(api_type, segment, context)
                    return (idx, result, None)
                except Exception as e:
                    return (idx, None, str(e))

            # 执行翻译循环
            if max_workers > 1:
                # 并发模式
                self.root.after(0, self.progress_text_var.set, f"正在并发翻译 (线程数: {max_workers})...")
                
                with ThreadPoolExecutor(max_workers=max_workers) as executor:
                    # 创建剩余任务
                    futures = {
                        executor.submit(process_segment, i): i 
                        for i in range(start_index, total_segments)
                    }
                    
                    completed_count = start_index
                    for future in as_completed(futures):
                        if not self.is_translating or self.paused_due_to_failures:
                            executor.shutdown(wait=False, cancel_futures=True)
                            break
                            
                        try:
                            idx, result, error = future.result()
                        except Exception as e:
                            idx = futures[future]
                            result = None
                            error = str(e)
                        
                        if result:
                            self.translated_segments[idx] = result
                            self.consecutive_failures = 0
                        else:
                            self.consecutive_failures += 1
                            self.translated_segments[idx] = f"[翻译错误: {error}]\n{self.source_segments[idx]}"
                            print(f"翻译段落 {idx + 1} 失败: {error}")
                            
                            if self.consecutive_failures >= self.max_consecutive_failures:
                                self.paused_due_to_failures = True
                                self.resume_from_index = idx  # 记录暂停位置（大概）
                        
                        completed_count += 1
                        progress = (completed_count / total_segments) * 100
                        self.root.after(0, self.progress_var.set, progress)
                        self.root.after(0, self.progress_text_var.set, f"正在翻译... {completed_count}/{total_segments} 段")
                        
                        # 定期保存和更新UI (不必每段都更新，减少开销)
                        if completed_count % 5 == 0:
                            self.save_progress_cache()
                            current_text = "\n\n".join(seg for seg in self.translated_segments if seg)
                            self.root.after(0, self.update_translated_text, current_text)
            else:
                # 单线程模式 (保持原逻辑以支持上下文)
                for idx in range(start_index, total_segments):
                    if not self.is_translating:
                        break
                    
                    # 重新调用 process_segment 逻辑
                    _, result, error = process_segment(idx)
                    
                    if result:
                        self.translated_segments[idx] = result
                        self.consecutive_failures = 0
                    else:
                        self.consecutive_failures += 1
                        self.translated_segments[idx] = f"[翻译错误: {error}]\n{self.source_segments[idx]}"
                        
                        if self.consecutive_failures >= self.max_consecutive_failures:
                            self.paused_due_to_failures = True
                            self.resume_from_index = idx
                            break
                    
                    progress = ((idx + 1) / total_segments) * 100
                    self.root.after(0, self.progress_var.set, progress)
                    self.root.after(0, self.progress_text_var.set, f"正在翻译... {idx + 1}/{total_segments} 段")
                    
                    # 实时更新
                    self.translated_text = "\n\n".join(self.translated_segments[:idx+1])
                    self.root.after(0, self.update_translated_text, self.translated_text)
                    self.save_progress_cache()
                    
                    time.sleep(0.2) # 避免单线程下的API限流

            # 翻译完成后的处理
            if self.is_translating and not self.paused_due_to_failures:
                # 最终更新一次完整文本
                self.translated_text = "\n\n".join(self.translated_segments)
                self.root.after(0, self.update_translated_text, self.translated_text)
                
                self.root.after(0, self.progress_text_var.set, "正在检查译文...")
                # 暂时只在单线程模式下重试，并发模式下重试逻辑较复杂
                if max_workers == 1:
                    self.verify_and_retry_segments(api_type)

                self.root.after(0, self.refresh_failed_segments_view)
                self.root.after(0, self.progress_var.set, 100)
                
                failed_count = sum(1 for s in self.translated_segments if s.startswith("[翻译错误") or s.startswith("[未翻译"))
                status_msg = (
                    f"翻译完成，有 {failed_count} 段可能需要检查"
                    if failed_count else "翻译完成!"
                )
                self.root.after(0, self.progress_text_var.set, status_msg)
                self.root.after(0, self.on_translation_complete)
                if failed_count == 0:
                    self.clear_progress_cache()
            else:
                status_msg = "翻译已停止"
                if self.paused_due_to_failures:
                    status_msg = "已暂停，等待API恢复后可继续"
                self.root.after(0, self.progress_text_var.set, status_msg)

        except Exception as e:
            self.root.after(
                0,
                messagebox.showerror,
                "错误",
                f"翻译过程中出错:\n{str(e)}"
            )
        finally:
            self.root.after(0, self.translate_btn.config, {'state': 'normal'})
            self.root.after(0, self.stop_btn.config, {'state': 'disabled'})
            self.is_translating = False

    def detect_language(self, text):
        """简单的语言检测：检查是否主要是中文"""
        if not text or len(text.strip()) == 0:
            return 'unknown'

        # 统计中文字符
        chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
        total_chars = len(re.findall(r'\S', text))

        if total_chars == 0:
            return 'unknown'

        chinese_ratio = chinese_chars / total_chars

        # 如果中文占比超过60%，认为是中文
        if chinese_ratio > 0.6:
            return 'zh'
        # 如果中文占比很低，可能是英文或其他语言
        elif chinese_ratio < 0.1:
            return 'en'
        else:
            return 'mixed'

    def translate_segment(self, api_type, text, context=None):
        """按当前API类型翻译单段文本（使用统一翻译引擎）"""
        target_language = self.get_target_language()
        target_is_chinese = self.is_target_language_chinese(target_language)
        target_is_english = self.is_target_language_english(target_language)

        # 检测语言，如果已经是目标语言就跳过翻译
        lang = self.detect_language(text)
        if (target_is_chinese and lang == 'zh') or (target_is_english and lang == 'en'):
            return text
        
        # 构建风格提示
        style = self.style_var.get()
        style_prompt_map = {
            "直译 (Literal)": "请进行精准直译，严格保留原文的句子结构和语气，不要过度意译。",
            "通俗小说 (Novel)": "请采用通俗小说的笔法，用词生动、流畅，注重情节的连贯性和人物语气的自然，符合目标语言读者的阅读习惯。",
            "学术专业 (Academic)": "请采用学术风格，用词严谨、专业，句式规范，确保术语准确，适合学术研究或专业人士阅读。",
            "武侠/古风 (Wuxia)": "请采用中国古典武侠或古风小说的笔触，用词典雅、古朴，半文半白，注重意境的渲染。",
            "新闻/媒体 (News)": "请采用新闻报道的风格，客观、简练、信息传达准确，符合新闻媒体的规范。"
        }
        style_guide = style_prompt_map.get(style, "")
        if style_guide:
            style_guide = f"风格要求：{style_guide}"
        
        # 调用翻译引擎
        # 注意：engine会自动处理翻译记忆、术语表、API调用、错误回退
        result = self.translation_engine.translate(
            text=text,
            target_lang=target_language,
            provider=self._engine_provider_name(api_type),
            use_memory=True,
            use_glossary=True,
            context=context,
            extra_prompt=style_guide
        )
        
        if result.success:
            return result.translated_text
        else:
            # 如果失败，抛出异常以便上层捕获处理（如记录失败段落）
            raise Exception(result.error or "未知翻译错误")



    def is_translation_incomplete(self, translated, source, target_language=None):
        """检测译文是否异常或未完成"""
        return review_is_translation_incomplete(
            translated,
            source,
            target_language or self.get_target_language(),
        )

    def verify_and_retry_segments(self, api_type):
        """翻译完成后检查并自动重试失败段落"""
        self.translated_segments, self.failed_segments = review_verify_and_retry_segments(
            self.source_segments,
            self.translated_segments,
            lambda source, idx: self.translate_segment(api_type, source),
            self.get_target_language(),
        )
        self.save_progress_cache()

    def refresh_failed_segments_view(self):
        """刷新失败段落列表和状态"""
        if not hasattr(self, 'failed_listbox'):
            return

        self.failed_listbox.delete(0, tk.END)
        self.selected_failed_index = None

        self.failed_source_text.config(state='normal')
        self.failed_source_text.delete('1.0', tk.END)
        self.failed_source_text.config(state='disabled')

        self.manual_translation_text.delete('1.0', tk.END)

        if not self.failed_segments:
            self.failed_status_var.set("暂无失败段落")
            return

        for item in self.failed_segments:
            snippet = item['source'].replace("\n", " ")
            if len(snippet) > 60:
                snippet = snippet[:60] + "..."
            self.failed_listbox.insert(tk.END, f"段 {item['index'] + 1}: {snippet}")

        self.failed_status_var.set(f"待处理段落: {len(self.failed_segments)} 个")

    def on_failed_select(self, event=None):
        """选中失败段落时展示详情"""
        if not self.failed_segments:
            return

        selection = self.failed_listbox.curselection()
        if not selection:
            return

        idx = selection[0]
        self.selected_failed_index = idx
        info = self.failed_segments[idx]

        self.failed_source_text.config(state='normal')
        self.failed_source_text.delete('1.0', tk.END)
        self.failed_source_text.insert('1.0', info['source'])
        self.failed_source_text.config(state='disabled')

        self.manual_translation_text.delete('1.0', tk.END)

    def retry_failed_segment(self):
        """对选中失败段落重新翻译"""
        if self.selected_failed_index is None or not self.failed_segments:
            messagebox.showinfo("提示", "请先选择需要重试的段落")
            return

        info = self.failed_segments[self.selected_failed_index]
        api_type = self.get_retry_api_type()
        if not self._ensure_provider_ready_or_prompt(api_type):
            return

        try:
            retry_text = self.translate_segment(api_type, info['source'])
        except Exception as e:
            messagebox.showerror("错误", f"重试翻译失败: {e}")
            return

        if self.is_translation_incomplete(retry_text, info['source'], target_language=self.get_target_language()):
            messagebox.showwarning("提示", "重试后仍未完成，请手动翻译。")
            return

        self.translated_segments[info['index']] = retry_text
        self.failed_segments.pop(self.selected_failed_index)
        self.rebuild_translated_text()
        self.refresh_failed_segments_view()
        self.save_progress_cache()
        messagebox.showinfo("成功", f"段 {info['index'] + 1} 已重新翻译并替换")

    def save_manual_translation(self):
        """将手动译文写回对应段落并保存到记忆库"""
        if self.selected_failed_index is None or not self.failed_segments:
            messagebox.showinfo("提示", "请先选择需要替换的段落")
            return

        manual_text = self.manual_translation_text.get('1.0', tk.END).strip()
        try:
            self.translated_segments, self.failed_segments, info = apply_manual_translation_review(
                self.translated_segments,
                self.failed_segments,
                self.selected_failed_index,
                manual_text,
            )
        except ValueError as exc:
            messagebox.showwarning("警告", str(exc))
            return

        source_text = info['source']
        
        # 2. 保存到翻译记忆库 (Linkage 1)
        try:
            target_lang = self.get_target_language()
            self.translation_memory.store(
                source_text=source_text,
                translated_text=manual_text,
                target_lang=target_lang,
                api_provider="manual",
                model="user_correction",
                quality_score=100  # 人工修正视为完美
            )
            print(f"已将人工修正存入记忆库: ID={info['index']}")
        except Exception as e:
            print(f"保存到记忆库失败: {e}")

        self.rebuild_translated_text()
        self.refresh_failed_segments_view()
        messagebox.showinfo("成功", f"段 {info['index'] + 1} 已使用手动译文替换并存入记忆库")
        self.save_progress_cache()

    def rebuild_translated_text(self):
        """根据分段译文重建完整译文"""
        self.translated_text = "\n\n".join(self.translated_segments) if self.translated_segments else ""
        self.update_translated_text(self.translated_text)



    def update_translated_text(self, text):
        """更新译文显示"""
        self.translated_text_widget.delete('1.0', tk.END)
        self.translated_text_widget.insert('1.0', text)
        # 自动滚动到底部
        self.translated_text_widget.see(tk.END)
        # 同步更新对照视图
        self.update_comparison_view()

    def on_translation_complete(self):
        """翻译完成后的处理"""
        # 刷新解析列表，以便用户可以进行解析
        self.refresh_analysis_listbox()
        # 确保对照视图也是最新的
        self.update_comparison_view()

        # 批量模式处理
        if self.is_batch_mode:
            # 自动导出
            self.auto_export_batch_file()
            # 标记当前任务完成
            for item in self.batch_queue:
                if item['status'] == 'processing':
                    item['status'] = 'done'
                    break
            self.save_batch_queue()
            # 继续下一个
            self.root.after(2000, self.process_next_batch_file)
            return

        if self.failed_segments:
            self.notebook.select(4)  # 切换到失败段落标签页 (索引: 0搜索, 1原文, 2译文, 3对照, 4失败)
            message = f"翻译完成，但 {len(self.failed_segments)} 个段落需要手动翻译或重试。"
            messagebox.showwarning("完成", message)
        else:
            self.notebook.select(1)  # 切换到译文标签页
            messagebox.showinfo("完成", "翻译已完成!")

    def auto_export_batch_file(self):
        """批量模式下的自动导出"""
        if not self.batch_output_dir or not self.translated_text:
            return
            
        try:
            target_language = self.get_target_language()
            safe_lang = re.sub(r'[\\/:*?"<>|]', "_", target_language).strip()
            original_path = Path(self.file_path_var.get())
            
            # 导出纯文本
            txt_name = f"{original_path.stem}_{safe_lang}译文.txt"
            txt_path = Path(self.batch_output_dir) / txt_name
            with open(txt_path, 'w', encoding='utf-8') as f:
                f.write(self.translated_text)
                
            print(f"批量导出成功: {txt_path}")
            
        except Exception as e:
            print(f"批量导出失败: {e}")
            # 标记为失败但继续
            for item in self.batch_queue:
                if item['status'] == 'processing':
                    item['status'] = 'failed'

    def export_translation(self):
        """导出翻译结果"""
        if not self.translated_text or not self.translated_text.strip():
            messagebox.showwarning("警告", "没有可导出的译文\n\n请先完成翻译后再导出")
            return

        # 建议默认文件名
        original_file = self.file_path_var.get()
        target_language = self.get_target_language()
        safe_lang = re.sub(r'[\\/:*?"<>|]', "_", target_language).strip() or "译文"
        
        # 扩展名处理
        ext = ".txt"
        file_types = [("文本文件", "*.txt")]
        
        # 如果是 DOCX 且处理器就绪，默认导出 DOCX
        if self.docx_handler and original_file.lower().endswith('.docx'):
            ext = ".docx"
            file_types = [("Word 文档", "*.docx"), ("文本文件", "*.txt")]

        if original_file:
            base_name = Path(original_file).stem
            default_name = f"{base_name}_{safe_lang}译文{ext}"
        else:
            default_name = f"{safe_lang}译文{ext}"

        filename = filedialog.asksaveasfilename(
            title="保存译文",
            defaultextension=ext,
            initialfile=default_name,
            filetypes=file_types
        )

        if filename:
            try:
                # 检查是否导出为 DOCX
                if filename.lower().endswith('.docx') and self.docx_handler:
                    self.docx_handler.save_translated_file(self.translated_segments, filename)
                    messagebox.showinfo("成功", f"格式保留的 Word 文档已保存到:\n{filename}")
                else:
                    with open(filename, 'w', encoding='utf-8') as f:
                        f.write(self.translated_text)
                    messagebox.showinfo("成功", f"译文已保存到:\n{filename}")

                # 完整导出后清除进度缓存
                if self.source_segments and len(self.translated_segments) == len(self.source_segments):
                    self.clear_progress_cache()
            except Exception as e:
                messagebox.showerror("错误", f"保存文件失败:\n{str(e)}")

    def export_audiobook(self):
        """导出有声书"""
        if not self.translated_text:
            messagebox.showwarning("警告", "没有可导出的译文")
            return
            
        ok, msg = self.audio_manager.check_dependency()
        if not ok:
            messagebox.showerror("错误", msg)
            return

        # 选择语音角色
        dialog = tk.Toplevel(self.root)
        dialog.title("导出有声书")
        dialog.geometry("400x200")
        
        ttk.Label(dialog, text="选择语音角色:").pack(pady=10)
        
        voices = self.audio_manager.get_voices()
        voice_var = tk.StringVar(value='zh-CN-XiaoxiaoNeural')
        voice_combo = ttk.Combobox(dialog, textvariable=voice_var, values=list(voices.keys()))
        voice_combo.pack(pady=5)
        
        # 显示友好的名称
        name_label = ttk.Label(dialog, text=voices[voice_var.get()])
        name_label.pack(pady=5)
        
        def on_voice_change(event):
            name_label.config(text=voices.get(voice_var.get(), ""))
        voice_combo.bind('<<ComboboxSelected>>', on_voice_change)
        
        def do_export():
            output_path = filedialog.asksaveasfilename(
                title="保存有声书",
                defaultextension=".mp3",
                filetypes=[("MP3 音频", "*.mp3")]
            )
            if not output_path:
                return
                
            dialog.destroy()
            
            # 后台生成
            def run_gen():
                self.progress_text_var.set("正在生成有声书 (这可能需要几分钟)...")
                try:
                    self.audio_manager.generate_audiobook(
                        self.translated_text[:100000], # 限制长度防止过长失败，实际应分段
                        output_path, 
                        voice_var.get()
                    )
                    messagebox.showinfo("成功", f"有声书已生成: {output_path}")
                except Exception as e:
                    messagebox.showerror("失败", f"生成失败: {e}")
                finally:
                    self.progress_text_var.set("就绪")
                    
            threading.Thread(target=run_gen).start()
            
        ttk.Button(dialog, text="开始生成", command=do_export).pack(pady=20)

    def update_text_display(self):
        """更新文本显示（预览或完整）"""
        if not self.current_text:
            return

        self.original_text.delete('1.0', tk.END)

        char_count = len(self.current_text)
        is_large_file = char_count > self.preview_limit

        if is_large_file and not self.show_full_text:
            # 显示预览
            preview_text = self.current_text[:self.preview_limit]
            preview_text += f"\n\n{'='*60}\n"
            preview_text += f"⚠️ 预览模式：仅显示前 {self.preview_limit:,} / {char_count:,} 字符\n"
            preview_text += f"点击上方'显示完整原文'按钮查看全文\n"
            preview_text += f"{'='*60}"
            self.original_text.insert('1.0', preview_text)
        else:
            # 显示完整文本
            self.original_text.insert('1.0', self.current_text)

    def toggle_full_text_display(self):
        """切换显示完整文本或预览"""
        if not self.current_text:
            return

        self.show_full_text = not self.show_full_text
        char_count = len(self.current_text)

        if self.show_full_text:
            # 切换到完整显示
            self.toggle_preview_btn.config(text="仅显示预览")
            self.file_info_var.set(f"✓ 显示完整文件 ({char_count:,} 字符)")
            self.progress_text_var.set("正在加载完整文本...")
            self.root.update()

            # 使用after延迟更新，避免界面冻结
            self.root.after(100, self._update_full_text)
        else:
            # 切换到预览
            self.toggle_preview_btn.config(text="显示完整原文")
            self.file_info_var.set(
                f"⚠️ 大文件 ({char_count:,} 字符) - 仅显示前 {self.preview_limit:,} 字符"
            )
            self.update_text_display()
            self.progress_text_var.set(f"已加载文件 | 字符数: {char_count:,}")

    def _update_full_text(self):
        """更新完整文本（在延迟后执行）"""
        self.update_text_display()
        char_count = len(self.current_text)
        word_count = len(self.current_text.split())
        self.progress_text_var.set(
            f"已加载完整文件 | 字符数: {char_count:,} | 词数: {word_count:,}"
        )

    def clear_all(self):
        """清空所有内容"""
        self.clear_all_internal(skip_ui_confirm=False)

    # ==================== 解析功能方法 ====================

    def refresh_analysis_listbox(self):
        """刷新解析标签页的段落列表"""
        self.analysis_listbox.delete(0, tk.END)

        if not self.translated_segments:
            self.analysis_status_var.set("翻译完成后可进行解析")
            return

        # 初始化解析结果列表（如果尚未初始化或长度不匹配）
        if len(self.analysis_segments) != len(self.translated_segments):
            self.analysis_segments = [''] * len(self.translated_segments)

        for i, seg in enumerate(self.translated_segments):
            # 显示段落编号和预览（前30字符）
            preview = seg[:30].replace('\n', ' ') + ('...' if len(seg) > 30 else '')
            status = "✓" if self.analysis_segments[i] else "○"
            self.analysis_listbox.insert(tk.END, f"{status} 段落 {i+1}: {preview}")

        analyzed_count = sum(1 for s in self.analysis_segments if s)
        total_count = len(self.translated_segments)
        self.analysis_status_var.set(f"已解析 {analyzed_count}/{total_count} 段")

    def on_analysis_segment_select(self, event=None):
        """当用户选择解析列表中的某一段时，显示对应内容"""
        selection = self.analysis_listbox.curselection()
        if not selection:
            return

        idx = selection[0]
        if idx >= len(self.translated_segments):
            return

        # 显示原文、译文和解析（如果有）
        source = self.source_segments[idx] if idx < len(self.source_segments) else ""
        translated = self.translated_segments[idx]
        analysis = self.analysis_segments[idx] if idx < len(self.analysis_segments) else ""

        self.analysis_text.delete('1.0', tk.END)
        self.analysis_text.insert(tk.END, f"【原文】\n{source}\n\n")
        self.analysis_text.insert(tk.END, f"【译文】\n{translated}\n\n")
        if analysis:
            self.analysis_text.insert(tk.END, f"【解析】\n{analysis}")
        else:
            self.analysis_text.insert(tk.END, '【解析】\n（尚未解析，点击"解析选中段落"按钮进行解析）')

    def analyze_selected_segment(self):
        """解析当前选中的单个段落"""
        selection = self.analysis_listbox.curselection()
        if not selection:
            messagebox.showwarning("警告", "请先在列表中选择一个段落")
            return

        idx = selection[0]
        if idx >= len(self.translated_segments):
            return

        # 检查是否正在解析
        if self.is_analyzing:
            messagebox.showinfo("提示", "正在解析中，请稍候...")
            return

        self.analysis_status_var.set(f"正在解析段落 {idx+1}...")

        # 在后台线程中执行解析
        def analyze_worker():
            try:
                source = self.source_segments[idx] if idx < len(self.source_segments) else ""
                translated = self.translated_segments[idx]

                result = self.call_api_for_analysis(source, translated)

                # 保存解析结果
                if idx < len(self.analysis_segments):
                    self.analysis_segments[idx] = result

                # 更新UI
                self.root.after(0, self._update_analysis_ui_after_single, idx, result)

            except Exception as e:
                self.root.after(0, lambda: messagebox.showerror("解析失败", f"解析段落 {idx+1} 时出错:\n{e}"))
                self.root.after(0, lambda: self.analysis_status_var.set("解析失败"))

        threading.Thread(target=analyze_worker, daemon=True).start()

    def _update_analysis_ui_after_single(self, idx, result):
        """单段解析完成后更新UI"""
        # 刷新列表显示状态
        self.refresh_analysis_listbox()
        # 重新选中该段落并显示解析结果
        self.analysis_listbox.selection_clear(0, tk.END)
        self.analysis_listbox.selection_set(idx)
        self.analysis_listbox.see(idx)
        self.on_analysis_segment_select()

    def start_batch_analysis(self):
        """开始批量解析所有已翻译段落"""
        if not self.translated_segments:
            messagebox.showwarning("警告", "请先完成翻译后再进行解析")
            return

        if self.is_analyzing:
            messagebox.showinfo("提示", "正在解析中，请稍候...")
            return

        if self.is_translating:
            messagebox.showwarning("警告", "请等待翻译完成后再进行解析")
            return

        # 确认是否覆盖已有解析
        existing_count = sum(1 for s in self.analysis_segments if s)
        if existing_count > 0:
            if not messagebox.askyesno("确认", f"已有 {existing_count} 段解析结果，是否全部重新解析？"):
                return

        # 初始化解析结果列表
        self.analysis_segments = [''] * len(self.translated_segments)
        self.is_analyzing = True

        # 更新按钮状态
        self.analyze_all_btn.config(state='disabled')
        self.stop_analysis_btn.config(state='normal')

        self.analysis_thread = threading.Thread(target=self._batch_analysis_worker, daemon=True)
        self.analysis_thread.start()

    def stop_analysis(self):
        """停止批量解析"""
        if self.is_analyzing:
            self.is_analyzing = False
            self.analysis_status_var.set("正在停止解析...")

    def _batch_analysis_worker(self):
        """批量解析后台工作线程"""
        total = len(self.translated_segments)
        success_count = 0
        fail_count = 0

        for i, translated in enumerate(self.translated_segments):
            if not self.is_analyzing:
                # 用户取消
                break

            source = self.source_segments[i] if i < len(self.source_segments) else ""

            self.root.after(0, lambda idx=i: self.analysis_status_var.set(f"正在解析 {idx+1}/{total}..."))
            self.root.after(0, lambda idx=i: self.progress_text_var.set(f"解析进度: {idx+1}/{total}"))

            try:
                result = self.call_api_for_analysis(source, translated)
                self.analysis_segments[i] = result
                success_count += 1
            except Exception as e:
                print(f"解析段落 {i+1} 失败: {e}")
                self.analysis_segments[i] = f"[解析失败: {e}]"
                fail_count += 1

            # 更新进度
            progress = (i + 1) / total * 100
            self.root.after(0, lambda p=progress: self.progress_var.set(p))

            # 每解析完一段刷新列表
            self.root.after(0, self.refresh_analysis_listbox)

            # 避免API限流
            time.sleep(0.5)

        was_cancelled = not self.is_analyzing
        self.is_analyzing = False

        # 恢复按钮状态
        self.root.after(0, lambda: self.analyze_all_btn.config(state='normal'))
        self.root.after(0, lambda: self.stop_analysis_btn.config(state='disabled'))

        self.root.after(0, lambda: self.progress_var.set(100))
        if was_cancelled:
            self.root.after(0, lambda: self.analysis_status_var.set(
                f"解析已停止。成功 {success_count} 段，失败 {fail_count} 段"
            ))
            self.root.after(0, lambda: self.progress_text_var.set("解析已停止"))
        else:
            self.root.after(0, lambda: self.analysis_status_var.set(
                f"解析完成！成功 {success_count} 段，失败 {fail_count} 段"
            ))
            self.root.after(0, lambda: self.progress_text_var.set("解析完成"))
        self.root.after(0, self.refresh_analysis_listbox)

    def call_api_for_analysis(self, source_text, translated_text):
        """调用API进行段落解析"""
        api_type = self.get_analysis_api_type()
        target_language = self.get_target_language()

        # 构建解析提示词
        prompt = f"""请对以下翻译内容进行详细解析和讲解。

【原文】
{source_text}

【译文】
{translated_text}

请从以下角度进行解析：
1. 内容概要：简要概括这段文字的主要内容
2. 关键信息：指出其中的关键概念、人物、事件或论点
3. 翻译说明：如有特殊术语或表达，说明翻译的处理方式
4. 延伸思考：提供相关的背景知识或思考角度

请用{target_language}回答。"""

        # 根据API类型调用对应的解析方法
        if api_type in self.custom_local_models:
            return self._analyze_with_custom_local_model(api_type, prompt)
        elif api_type == 'gemini':
            return self._analyze_with_gemini(prompt)
        elif api_type == 'openai':
            return self._analyze_with_openai(prompt)
        elif api_type == 'custom':
            return self._analyze_with_custom_api(prompt)
        elif api_type == 'lm_studio':
            return self._analyze_with_lm_studio(prompt)
        else:
            raise ValueError(f"不支持的解析API类型: {api_type}")

    def generate_text_with_selected_api(self, prompt, preferred_api_type=None):
        """使用当前选中的解析 API 进行通用文本生成。"""
        api_type = preferred_api_type or self.get_analysis_api_type()

        if api_type in self.custom_local_models:
            return self._analyze_with_custom_local_model(api_type, prompt)
        elif api_type == 'gemini':
            return self._analyze_with_gemini(prompt)
        elif api_type == 'openai':
            return self._analyze_with_openai(prompt)
        elif api_type == 'custom':
            return self._analyze_with_custom_api(prompt)
        elif api_type == 'lm_studio':
            return self._analyze_with_lm_studio(prompt)
        else:
            raise ValueError(f"不支持的文本生成 API 类型: {api_type}")

    def _analyze_with_custom_local_model(self, model_key, prompt):
        """使用自定义本地模型进行解析"""
        if not OPENAI_SUPPORT:
            raise ImportError("缺少 openai 库，无法调用本地模型")

        if model_key not in self.custom_local_models:
            raise ValueError(f"本地模型 '{model_key}' 未配置")

        config = self.custom_local_models[model_key]

        client = openai.OpenAI(
            api_key=config.get('api_key') or 'lm-studio',
            base_url=config['base_url']
        )

        response = client.chat.completions.create(
            model=config['model_id'],
            messages=[
                {"role": "system", "content": "你是一个专业的文本分析助手，擅长对翻译内容进行深度解析和讲解。"},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3
        )

        return response.choices[0].message.content

    def _analyze_with_gemini(self, prompt):
        """使用Gemini API进行解析"""
        if not GEMINI_SUPPORT:
            raise ImportError("缺少 google-generativeai 库")

        api_key = self.api_configs['gemini'].get('api_key', '')
        model_name = self.api_configs['gemini'].get('model', 'gemini-2.5-flash')

        if not api_key:
            raise ValueError("未配置 Gemini API Key")

        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(model_name)
        response = model.generate_content(prompt)

        return response.text

    def _analyze_with_openai(self, prompt):
        """使用OpenAI API进行解析"""
        if not OPENAI_SUPPORT:
            raise ImportError("缺少 openai 库")

        api_key = self.api_configs['openai'].get('api_key', '')
        model_name = self.api_configs['openai'].get('model', 'gpt-3.5-turbo')
        base_url = self.api_configs['openai'].get('base_url', '')

        if not api_key:
            raise ValueError("未配置 OpenAI API Key")

        client_kwargs = {'api_key': api_key}
        if base_url:
            client_kwargs['base_url'] = base_url

        client = openai.OpenAI(**client_kwargs)
        response = client.chat.completions.create(
            model=model_name,
            messages=[
                {"role": "system", "content": "你是一个专业的文本分析助手，擅长对翻译内容进行深度解析和讲解。"},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3
        )

        return response.choices[0].message.content

    def _analyze_with_custom_api(self, prompt):
        """使用自定义API进行解析"""
        if not REQUESTS_SUPPORT:
            raise ImportError("缺少 requests 库")

        api_key = self.api_configs['custom'].get('api_key', '')
        model_name = self.api_configs['custom'].get('model', '')
        base_url = self.api_configs['custom'].get('base_url', '')

        if not base_url:
            raise ValueError("未配置自定义API地址")

        headers = {"Content-Type": "application/json"}
        if api_key:
            headers["Authorization"] = f"Bearer {api_key}"

        data = {
            "model": model_name,
            "messages": [
                {"role": "system", "content": "你是一个专业的文本分析助手，擅长对翻译内容进行深度解析和讲解。"},
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.3
        }

        response = requests.post(
            f"{base_url.rstrip('/')}/chat/completions",
            headers=headers,
            json=data,
            timeout=120
        )
        response.raise_for_status()

        result = response.json()
        return result['choices'][0]['message']['content']

    def _analyze_with_lm_studio(self, prompt):
        """使用LM Studio本地模型进行解析"""
        if not OPENAI_SUPPORT:
            raise ImportError("缺少 openai 库")

        config = self.api_configs.get('lm_studio', DEFAULT_LM_STUDIO_CONFIG)

        client = openai.OpenAI(
            api_key=config.get('api_key') or 'lm-studio',
            base_url=config.get('base_url', 'http://127.0.0.1:1234/v1')
        )

        response = client.chat.completions.create(
            model=config.get('model', 'qwen2.5-7b-instruct-1m'),
            messages=[
                {"role": "system", "content": "你是一个专业的文本分析助手，擅长对翻译内容进行深度解析和讲解。"},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3
        )

        return response.choices[0].message.content

    def copy_analysis_content(self):
        """复制解析内容到剪贴板"""
        content = self.analysis_text.get('1.0', tk.END).strip()
        if not content:
            messagebox.showinfo("提示", "没有可复制的内容")
            return

        self.root.clipboard_clear()
        self.root.clipboard_append(content)
        messagebox.showinfo("成功", "已复制到剪贴板")

    def export_analysis(self):
        """导出解析结果"""
        if not self.analysis_segments or not any(self.analysis_segments):
            messagebox.showwarning("警告", "没有可导出的解析内容\n请先完成段落解析")
            return

        # 生成默认文件名
        original_file = self.file_path_var.get()
        if original_file:
            base_name = Path(original_file).stem
            default_name = f"{base_name}_解析.txt"
        else:
            default_name = "解析结果.txt"

        file_path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("文本文件", "*.txt"), ("所有文件", "*.*")],
            initialfile=default_name,
            title="导出解析结果"
        )

        if not file_path:
            return

        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write("=" * 60 + "\n")
                f.write("书籍翻译工具 - 段落解析结果\n")
                f.write("=" * 60 + "\n\n")

                for i, (source, translated, analysis) in enumerate(zip(
                    self.source_segments,
                    self.translated_segments,
                    self.analysis_segments
                )):
                    f.write(f"{'='*40}\n")
                    f.write(f"段落 {i+1}\n")
                    f.write(f"{'='*40}\n\n")
                    f.write(f"【原文】\n{source}\n\n")
                    f.write(f"【译文】\n{translated}\n\n")
                    f.write(f"【解析】\n{analysis if analysis else '（未解析）'}\n\n")

            analyzed_count = sum(1 for s in self.analysis_segments if s)
            total_count = len(self.analysis_segments)
            messagebox.showinfo(
                "导出成功",
                f"解析结果已保存到:\n{file_path}\n\n"
                f"共 {total_count} 段，已解析 {analyzed_count} 段"
            )

        except Exception as e:
            messagebox.showerror("导出失败", f"保存文件时出错:\n{e}")

    def export_bilingual_epub(self):
        """导出双语对照 EPUB 电子书"""
        if not self.translated_segments:
            messagebox.showwarning("警告", "没有可导出的译文")
            return

        # 检查是否安装了 ebooklib
        if not EPUB_SUPPORT:
            messagebox.showerror("错误", "未安装 ebooklib 库，无法导出 EPUB。\n请运行 py -m pip install ebooklib")
            return

        # 选择保存路径
        original_file = self.file_path_var.get()
        default_name = f"{Path(original_file).stem}_双语版.epub" if original_file else "双语书籍.epub"
        
        file_path = filedialog.asksaveasfilename(
            defaultextension=".epub",
            filetypes=[("EPUB 电子书", "*.epub")],
            initialfile=default_name,
            title="导出双语 EPUB"
        )
        
        if not file_path:
            return

        try:
            # 创建 EPUB 书籍
            book = epub.EpubBook()
            
            # 设置元数据
            title = Path(original_file).stem if original_file else "Translation"
            book.set_identifier(str(uuid.uuid4()))
            book.set_title(f"{title} (双语版)")
            book.set_language(self.get_target_language())
            book.add_author("Book Translator AI")

            # 创建章节
            chapters = []
            # 将每 50 段作为一个章节，避免单页过长
            chunk_size = 50
            total_segments = len(self.source_segments)
            
            # 简单的 CSS 样式
            style = """
                body { font-family: sans-serif; }
                .segment { margin-bottom: 1.5em; border-bottom: 1px dashed #ccc; padding-bottom: 1em; }
                .original { color: #666; font-size: 0.9em; margin-bottom: 0.5em; }
                .translation { color: #000; font-size: 1.0em; font-weight: bold; }
            """
            css_item = epub.EpubItem(uid="style_nav", file_name="style/nav.css", media_type="text/css", content=style)
            book.add_item(css_item)

            for i in range(0, total_segments, chunk_size):
                chunk_source = self.source_segments[i:i+chunk_size]
                chunk_trans = self.translated_segments[i:i+chunk_size]
                
                # 确保长度一致
                if len(chunk_trans) < len(chunk_source):
                    chunk_trans.extend([""] * (len(chunk_source) - len(chunk_trans)))

                # 构建 HTML 内容
                html_content = ["<h1>第 {} 部分</h1>".format(i // chunk_size + 1)]
                for src, trans in zip(chunk_source, chunk_trans):
                    # 处理换行符
                    src_html = src.replace("\n", "<br/>")
                    trans_html = trans.replace("\n", "<br/>")
                    
                    html_content.append(f"""
                    <div class="segment">
                        <div class="original">{src_html}</div>
                        <div class="translation">{trans_html}</div>
                    </div>
                    """)

                chapter = epub.EpubHtml(title=f"Part {i // chunk_size + 1}", file_name=f"part_{i // chunk_size + 1}.xhtml", lang='zh')
                chapter.content = "".join(html_content)
                chapter.add_item(css_item)
                
                book.add_item(chapter)
                chapters.append(chapter)

            # 定义书籍骨架
            book.toc = tuple(chapters)
            book.add_item(epub.EpubNcx())
            book.add_item(epub.EpubNav())
            book.spine = ['nav'] + chapters

            # 保存文件
            epub.write_epub(file_path, book, {})
            
            messagebox.showinfo("成功", f"双语 EPUB 已导出:\n{file_path}")

        except Exception as e:
            messagebox.showerror("导出失败", f"生成 EPUB 时出错:\n{e}")

    # --- 在线搜索相关方法 ---
    def setup_search_tab(self):
        """设置在线书城：包含 '全网搜索' 和 '社区图书馆' 两个子标签"""
        self.library_panel = LibraryPanel(
            self.library_frame,
            on_sidebar_search_click=self.on_sidebar_search_click,
            on_search_click=self.on_search_click,
            on_ai_search_click=self.on_ai_search_click,
            on_open_online_config=self.open_online_config,
            on_search_result_select=self.on_search_result_select,
            on_prev_page=self.on_prev_page,
            on_next_page=self.on_next_page,
            on_page_slider_release=self.on_page_slider_release,
            on_auto_categorize_click=self.on_auto_categorize_click,
            on_download_click=self.on_download_click,
            on_refresh_community_list=self.refresh_community_list,
            on_open_community_upload=self.open_community_upload,
            on_open_admin_audit=self.open_admin_audit,
            on_download_community_book=self.download_community_book,
            on_copy_community_link=self.copy_community_link,
        )

        self.library_notebook = self.library_panel.library_notebook
        self.comm_status_var = self.library_panel.comm_status_var
        self.comm_tree = self.library_panel.comm_tree
        self.comm_menu = self.library_panel.comm_menu
        self.admin_btn = self.library_panel.admin_btn
        self.cat_tree = self.library_panel.cat_tree
        self.lang_vars = self.library_panel.lang_vars
        self.search_frame = self.library_panel.search_frame
        self.search_query_var = self.library_panel.search_query_var
        self.search_entry = self.library_panel.search_entry
        self.search_source_var = self.library_panel.search_source_var
        self.search_tree = self.library_panel.search_tree
        self.prev_btn = self.library_panel.prev_btn
        self.page_label_var = self.library_panel.page_label_var
        self.page_slider = self.library_panel.page_slider
        self.next_btn = self.library_panel.next_btn
        self.search_detail_var = self.library_panel.search_detail_var
        self.download_btn = self.library_panel.download_btn
        self.search_status_var = self.library_panel.search_status_var

        self.current_search_results = []
        self.selected_result = None
        self.current_page = 1

        self.refresh_community_list()

    def refresh_community_list(self):
        """刷新社区书籍列表"""
        self.comm_tree.delete(*self.comm_tree.get_children())
        try:
            books = self.community_manager.get_public_books()
            for b in books:
                self.comm_tree.insert("", "end", values=(
                    b.get('title', 'Unknown'),
                    b.get('author', 'Unknown'),
                    b.get('description', ''),
                    b.get('uploader', 'Anon'),
                    b.get('size', ''),
                    b.get('date', '')
                ), tags=(b['url'],)) # Store URL in tag
        except Exception as e:
            messagebox.showerror("刷新失败", str(e))

    def download_community_book(self):
        """下载社区书籍"""
        selected = self.comm_tree.selection()
        if not selected: return
        item = self.comm_tree.item(selected[0])
        url = self.comm_tree.item(selected[0], "tags")[0]
        title = item['values'][0]
        
        if messagebox.askyesno("下载", f"确定下载书籍: {title}?"):
            try:
                # Reuse web importer logic or simple download
                self.progress_text_var.set(f"正在下载: {title}...")
                
                def run():
                    try:
                        resp = requests.get(url)
                        if resp.status_code == 200:
                            # Save to temp
                            ext = ".pdf" if url.endswith(".pdf") else ".txt" # Simple heuristic
                            if "epub" in url: ext = ".epub"
                            
                            path = Path("downloads") / f"{title}{ext}"
                            path.parent.mkdir(exist_ok=True)
                            
                            with open(path, 'wb') as f:
                                f.write(resp.content)
                                
                            self.root.after(0, lambda: self._load_downloaded_book(str(path)))
                        else:
                            raise Exception(f"Download failed: {resp.status_code}")
                    except Exception as e:
                        self.root.after(0, lambda: messagebox.showerror("错误", str(e)))
                
                threading.Thread(target=run, daemon=True).start()
                
            except Exception as e:
                messagebox.showerror("错误", str(e))

    def _load_downloaded_book(self, path):
        if messagebox.askyesno("下载完成", "下载成功！是否立即加载翻译？"):
            self.file_path_var.set(path)
            self.load_file_content(path)
            # Switch to Workstation tab
            self.main_notebook.select(0)

    def copy_community_link(self):
        selected = self.comm_tree.selection()
        if not selected: return
        url = self.comm_tree.item(selected[0], "tags")[0]
        self.root.clipboard_clear()
        self.root.clipboard_append(url)

    def open_community_upload(self):
        """打开上传分享对话框"""
        dialog = tk.Toplevel(self.root)
        dialog.title("上传书籍到社区")
        dialog.geometry("500x450")
        
        frame = ttk.Frame(dialog, padding=20)
        frame.pack(fill=tk.BOTH, expand=True)
        
        ttk.Label(frame, text="1. 选择文件").pack(anchor=tk.W)
        file_var = tk.StringVar()
        f_entry = ttk.Entry(frame, textvariable=file_var)
        f_entry.pack(fill=tk.X)
        ttk.Button(frame, text="浏览", command=lambda: file_var.set(filedialog.askopenfilename())).pack(anchor=tk.E, pady=2)
        
        ttk.Label(frame, text="2. 书籍信息").pack(anchor=tk.W, pady=(10, 0))
        
        ttk.Label(frame, text="标题:").pack(anchor=tk.W)
        title_var = tk.StringVar()
        ttk.Entry(frame, textvariable=title_var).pack(fill=tk.X)
        
        ttk.Label(frame, text="作者:").pack(anchor=tk.W)
        author_var = tk.StringVar()
        ttk.Entry(frame, textvariable=author_var).pack(fill=tk.X)
        
        ttk.Label(frame, text="简介:").pack(anchor=tk.W)
        desc_var = tk.StringVar()
        ttk.Entry(frame, textvariable=desc_var).pack(fill=tk.X)
        
        ttk.Label(frame, text="上传者昵称:").pack(anchor=tk.W)
        user_var = tk.StringVar(value="Anonymous")
        ttk.Entry(frame, textvariable=user_var).pack(fill=tk.X)
        
        # === AI 自动识别模块 ===
        def auto_fill_metadata():
            path = file_var.get()
            if not path or not os.path.exists(path):
                messagebox.showerror("错误", "请先选择有效文件")
                return
                
            ai_btn.config(state='disabled', text="AI 分析中...")
            self.root.update()
            
            def run_ai():
                try:
                    # 1. 读取文件前 8000 字符
                    content = self.file_processor.read_file(path)
                    if not content: raise Exception("无法读取文件内容")
                    preview_text = content[:8000]
                    
                    # 2. 构建 Prompt
                    prompt = (
                        "你是一个专业的图书管理员。请分析以下书籍片段，并提取元数据。\n"
                        "请严格返回 JSON 格式，包含以下字段：\n"
                        "- title: 书名 (如果无法确定，根据内容拟定)\n"
                        "- author: 作者 (如果未知，填 'Unknown')\n"
                        "- description: 200字以内的精彩简介，包含核心主题、风格和亮点。\n"
                        "\n"
                        f"书籍片段：\n{preview_text}..."
                    )
                    # 3. 调用当前选中的文本生成 API
                    api_type = self.get_analysis_api_type()
                    if not self._provider_ready_for_gui(api_type):
                        api_type = self.get_translation_api_type()

                    strict_json_prompt = (
                        "你是一个专业的图书管理员。请严格返回 JSON 对象，不要输出 Markdown 代码块，也不要附加解释。\n\n"
                        + prompt
                    )
                    response = self.generate_text_with_selected_api(strict_json_prompt, preferred_api_type=api_type)
                    
                    # 4. 解析 JSON
                    # 简单的 JSON 提取逻辑
                    try:
                        match = re.search(r'\{.*\}', response, re.DOTALL)
                        if not match:
                            raise ValueError("AI 未返回有效 JSON")
                        json_str = match.group(0)
                        data = json.loads(json_str)
                        
                        # 回到主线程更新 UI
                        self.root.after(0, lambda: update_ui(data))
                    except:
                        # 如果 AI 没返回标准 JSON，尝试简单的文本提取或报错
                        print(f"AI Response (Raw): {response}")
                        self.root.after(0, lambda: fail("AI 返回格式难以解析，请重试或手动填写。"))
                        
                except Exception as e:
                    self.root.after(0, lambda: fail(str(e)))

            def update_ui(data):
                if data.get('title'): title_var.set(data['title'])
                if data.get('author'): author_var.set(data['author'])
                if data.get('description'): desc_var.set(data['description'])
                ai_btn.config(state='normal', text="✨ AI 智能识别 (重新生成)")
                messagebox.showinfo("成功", "AI 已为您自动填写书籍信息！")

            def fail(msg):
                messagebox.showerror("AI 分析失败", msg)
                ai_btn.config(state='normal', text="✨ AI 智能识别")

            threading.Thread(target=run_ai, daemon=True).start()

        ai_btn = ttk.Button(frame, text="✨ AI 智能识别 (Auto-Fill)", command=auto_fill_metadata)
        ai_btn.pack(pady=5, fill=tk.X)
        # =======================
        
        def do_upload():
            path = file_var.get()
            if not path or not os.path.exists(path):
                messagebox.showerror("错误", "请选择有效文件")
                return
            
            btn.config(state='disabled', text="上传中...")
            
            def run():
                try:
                    self.community_manager.submit_book(
                        path,
                        title_var.get() or Path(path).stem,
                        author_var.get() or "Unknown",
                        desc_var.get(),
                        user_var.get()
                    )
                    self.root.after(0, lambda: success())
                except Exception as e:
                    self.root.after(0, lambda: fail(str(e)))
            
            threading.Thread(target=run, daemon=True).start()
            
        def success():
            messagebox.showinfo("提交成功", "书籍已上传并发布到社区！")
            dialog.destroy()
            self.refresh_community_list() # 自动刷新列表
            
        def fail(msg):
            messagebox.showerror("失败", msg)
            btn.config(state='normal', text="立即分享")

        btn = ttk.Button(frame, text="立即分享 (公开)", command=do_upload)
        btn.pack(pady=20, fill=tk.X)


    def open_admin_audit(self):
        """打开管理员管理界面（用于删除记录）"""
        expected_password = (self.config_manager.get_admin_password() or "").strip()
        if not expected_password:
            messagebox.showwarning(
                "未启用",
                "图书馆管理功能未启用。\n请通过环境变量 BOOK_TRANSLATOR_ADMIN_PASSWORD 或配置项 security.admin_password 设置管理员密码。"
            )
            return

        pwd = simpledialog.askstring("管理员登录", "请输入管理员密码:", show="*")
        if pwd is None:
            return
        if not hmac.compare_digest(pwd, expected_password):
            messagebox.showerror("错误", "密码错误")
            return

        audit_win = tk.Toplevel(self.root)
        audit_win.title("图书馆管理 (Library Admin)")
        audit_win.geometry("800x500")

        columns = ("id", "title", "uploader", "date", "status")
        tree = ttk.Treeview(audit_win, columns=columns, show="headings")
        tree.heading("id", text="ID")
        tree.heading("title", text="标题")
        tree.heading("uploader", text="上传者")
        tree.heading("date", text="日期")
        tree.heading("status", text="状态")
        tree.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        def refresh():
            tree.delete(*tree.get_children())
            books = self.community_manager.get_public_books()
            for b in books:
                tree.insert(
                    "",
                    "end",
                    values=(
                        b.get("id", ""),
                        b.get("title", ""),
                        b.get("uploader", ""),
                        b.get("date", ""),
                        b.get("status", "approved"),
                    ),
                )

        refresh()

        btn_frame = ttk.Frame(audit_win, padding=10)
        btn_frame.pack(fill=tk.X)

        def delete_entry():
            sel = tree.selection()
            if not sel:
                return
            bid = str(tree.item(sel[0])["values"][0])

            if messagebox.askyesno("确认", "确定从公共列表中删除此书籍吗？"):
                library = self.community_manager.get_public_books()
                new_library = [b for b in library if str(b.get("id", "")) != bid]
                self.community_manager._save_json(self.community_manager.library_file, new_library)
                messagebox.showinfo("成功", "已删除")
                refresh()
                self.refresh_community_list()

        ttk.Button(btn_frame, text="🗑️ 删除选中条目", command=delete_entry).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="刷新", command=refresh).pack(side=tk.RIGHT, padx=5)

    def on_prev_page(self):
        if self.current_page > 1:
            self.current_page -= 1
            self.perform_paged_search()

    def on_next_page(self):
        self.current_page += 1
        self.perform_paged_search()

    def on_page_slider_release(self, event):
        """滑块释放时跳转"""
        new_page = self.page_slider.get()
        if new_page != self.current_page:
            self.current_page = new_page
            self.perform_paged_search()

    def perform_paged_search(self):
        """执行分页搜索"""
        query = self.search_query_var.get().strip()
        source = self.search_source_var.get()
        if not query: return
        
        # 同步控件状态
        self.page_label_var.set(f"第 {self.current_page} 页")
        self.page_slider.set(self.current_page)
        
        self.search_status_var.set(f"正在搜索第 {self.current_page} 页: {query}...")
        
        # 更新按钮状态
        self.prev_btn.config(state='normal' if self.current_page > 1 else 'disabled')
        
        # 开启线程搜索
        threading.Thread(target=self.perform_search, args=(query, source, self.current_page), daemon=True).start()

    def on_sidebar_search_click(self):
        """处理侧边栏组合搜索"""
        # 1. 收集选中的分类
        selected_cats = []
        for item_id in self.cat_tree.selection():
            item_text = self.cat_tree.item(item_id, "text")
            # 提取英文部分作为关键词 (如果有)
            if "(" in item_text:
                keyword = item_text.split("(")[1].strip(")")
                selected_cats.append(keyword)
            else:
                selected_cats.append(item_text)
        
        # 2. 收集选中的语言
        selected_langs = [lang for lang, var in self.lang_vars.items() if var.get()]
        
        # 3. 收集输入框的额外关键词
        extra_keywords = self.search_query_var.get().strip()
        
        # 4. 组合查询语句
        query_parts = []
        if selected_cats:
            query_parts.extend(selected_cats)
        if selected_langs:
            query_parts.extend(selected_langs)
        if extra_keywords:
            query_parts.append(extra_keywords)
            
        final_query = " ".join(query_parts)
        
        if not final_query:
            messagebox.showwarning("提示", "请至少选择一个分类、语言或输入关键词")
            return
            
        # 设置到输入框并触发搜索
        self.search_query_var.set(final_query)
        self.on_search_click()

    def on_category_select(self, event):
        """(保留单机行为，但多选模式下主要由按钮触发)"""
        pass # 多选模式下，点击不直接触发，由按钮触发

    def on_random_browse_click(self):
        """随机浏览"""
        import random
        # 常见分类和热门关键词
        keywords = [
            "Best sci-fi books 2024", "Classic literature", "Python programming", 
            "History of China", "Modern philosophy", "Psychology bestsellers",
            "Artificial Intelligence", "Machine Learning", "Space exploration",
            "Fantasy novels", "Mystery thrillers", "Biographies",
            "Cooking and Food", "Self-help", "Finance and Investment"
        ]
        random_keyword = random.choice(keywords)
        self.search_query_var.set(random_keyword)
        self.on_search_click()

    def on_auto_categorize_click(self):
        """获取网站原本的分类信息（支持单选、文件夹批量、全量批量）"""
        selection = self.search_tree.selection()
        items_to_process = []
        
        # 1. 如果没有选中任何项 -> 询问是否处理全部
        if not selection:
            if not self.current_search_results:
                messagebox.showwarning("提示", "列表为空，请先搜索")
                return
            
            if messagebox.askyesno("批量获取", f"您未选中任何书籍。\n是否要对当前列表中的 {len(self.current_search_results)} 本书全部获取详细分类？\n(这可能需要一些时间)"):
                items_to_process = self.current_search_results
            else:
                return

        # 2. 如果选中了某项
        else:
            item_id = selection[0]
            
            # 情况 A: 选中了分类文件夹 -> 处理该分类下的所有书籍
            if "_cat_" in item_id:
                cat_name = self.search_tree.item(item_id, "values")[0].replace("📂 ", "")
                # 找到所有属于该分类的书籍
                items_to_process = [res for res in self.current_search_results if res.get('category') == cat_name]
                
                if not items_to_process: # Fallback just in case
                    return
                    
                if not messagebox.askyesno("批量获取", f"是否获取 '{cat_name}' 分类下 {len(items_to_process)} 本书的详细分类？"):
                    return
            
            # 情况 B: 选中了具体的书籍 -> 处理单本书
            else:
                item_values = self.search_tree.item(item_id, "values")
                title = item_values[0]
                for res in self.current_search_results:
                    if res.get('title') == title:
                        items_to_process = [res]
                        break
        
        if not items_to_process:
            return

        self.search_status_var.set(f"准备获取 {len(items_to_process)} 本书的分类信息...")
        self.download_btn.config(state="disabled") # 暂时禁用下载防止冲突
        
        def run_batch_fetch():
            count = 0
            total = len(items_to_process)
            
            for i, book_info in enumerate(items_to_process):
                title = book_info.get('title', 'Unknown')
                url = book_info.get('url', '')
                source = book_info.get('source', '')
                
                if not url: continue
                
                self.root.after(0, lambda t=title, c=count+1, tot=total: self.search_status_var.set(f"[{c}/{tot}] 获取分类: {t[:20]}..."))
                
                try:
                    # 获取详细分类
                    new_category = self.online_search_manager.get_book_category(url, source)
                    
                    if new_category and new_category != "Unknown":
                        book_info['category'] = new_category # 更新数据
                    
                    # 稍微延时避免请求过快
                    if total > 1:
                        time.sleep(0.5)
                        
                except Exception as e:
                    print(f"Error fetching cat for {title}: {e}")
                
                count += 1
            
            # 全部完成后刷新界面
            self.root.after(0, lambda: self._refresh_search_tree_grouped())
            self.root.after(0, lambda: self.search_status_var.set(f"批量获取完成，已更新 {count} 本书的分类"))
            self.root.after(0, lambda: self.download_btn.config(state="normal"))

        threading.Thread(target=run_batch_fetch, daemon=True).start()

    def _refresh_search_tree_grouped(self):
        """重新根据当前数据刷新树状列表"""
        # 清空当前树
        for item in self.search_tree.get_children():
            self.search_tree.delete(item)
            
        # 重新分组
        categories = {}
        for res in self.current_search_results:
            cat = res.get('category', 'Uncategorized')
            if cat not in categories:
                categories[cat] = []
            categories[cat].append(res)
        
        # 重新插入
        for i, (cat, cat_books) in enumerate(categories.items()):
            cat_id = f"group_{i}"
            group_source = cat_books[0].get('source', '') if cat_books else ''
            # 文件夹节点
            self.search_tree.insert(
                "",
                tk.END,
                iid=cat_id,
                values=("📂 " + cat, "", "", "", "", group_source, "分类"),
                open=True
            )
            
            for j, res in enumerate(cat_books):
                # 书籍节点
                self.search_tree.insert(cat_id, tk.END, iid=f"book_{i}_{j}", values=(
                    res.get('title', '未知'),
                    res.get('author', '未知'),
                    res.get('language', ''),
                    res.get('extension', ''),
                    res.get('size', ''),
                    res.get('source', res.get('source', '')),
                    cat
                ))


    def on_ai_search_click(self):
        """AI 智能寻书点击处理"""
        query = self.search_query_var.get().strip()
        if not query:
            messagebox.showwarning("警告", "请输入您的寻书需求（例如：'找一本关于Python数据分析的畅销书'）")
            return
            
        source = self.search_source_var.get()
        self.search_status_var.set(f"AI 正在分析需求: {query}...")
        self.download_btn.config(state="disabled")
        
        # 清空当前列表
        for item in self.search_tree.get_children():
            self.search_tree.delete(item)
            
        # 开启线程搜索
        threading.Thread(target=self.perform_ai_search, args=(query, source), daemon=True).start()

    def perform_ai_search(self, query, source):
        """执行 AI 寻书逻辑"""
        try:
            def callback(msg):
                self.root.after(0, lambda: self.search_status_var.set(msg))
            
            # 1. 调用 BookHunter
            results = self.book_hunter.hunt(query, source, callback=callback)
            
            # 2. (可选) AI 再次筛选
            # callback("🤖 AI 正在筛选最佳匹配...")
            # filtered_results = self.book_hunter.ai_filter_results(query, results)
            # results = filtered_results if filtered_results else results
            
            self.current_search_results = results
            
            def update_ui():
                # 清空旧结果
                for item in self.search_tree.get_children():
                    self.search_tree.delete(item)

                if not results:
                    self.search_status_var.set("AI 未找到相关书籍")
                    messagebox.showinfo("提示", "AI 分析了您的需求，但未找到匹配的书籍。")
                    return
                
                # 按分类分组
                categories = {}
                for res in results:
                    cat = res.get('category', 'Uncategorized')
                    if cat not in categories:
                        categories[cat] = []
                    categories[cat].append(res)
                
                # 清空并以树状展示
                for i, (cat, cat_books) in enumerate(categories.items()):
                    # 插入分类父节点
                    cat_id = f"ai_cat_{i}"
                    self.search_tree.insert("", tk.END, iid=cat_id, values=("📂 " + cat, "", "", "", "", source, "分类"), open=True)
                    
                    for j, res in enumerate(cat_books):
                        self.search_tree.insert(cat_id, tk.END, iid=f"ai_book_{i}_{j}", values=(
                            res.get('title', '未知'),
                            res.get('author', '未知'),
                            res.get('language', ''),
                            res.get('extension', ''),
                            res.get('size', ''),
                            res.get('source', source),
                            cat
                        ))
                
                self.search_status_var.set(f"AI 寻书完成，找到 {len(results)} 本推荐书籍")
                
            self.root.after(0, update_ui)
            
        except Exception as e:
            self.root.after(0, lambda: messagebox.showerror("AI 寻书错误", str(e)))
            self.root.after(0, lambda: self.search_status_var.set("寻书出错"))

    def on_search_click(self):
        """点击搜索按钮"""
        query = self.search_query_var.get().strip()
        if not query:
            messagebox.showwarning("警告", "请输入搜索关键词")
            return
            
        source = self.search_source_var.get()
        self.search_status_var.set(f"正在从 {source} 搜索: {query}...")
        self.download_btn.config(state="disabled")
        
        # 重置分页
        self.current_page = 1
        self.perform_paged_search()

    def perform_search(self, query, source, page=1):
        """执行搜索逻辑（后台线程）"""
        try:
            if source == "Z-Library":
                results = self.online_search_manager.search_zlibrary(query, page)
            else:
                results = self.online_search_manager.search_annas_archive(query, page)
                
            self.current_search_results = results
            
            def update_ui():
                # 清空旧结果
                for item in self.search_tree.get_children():
                    self.search_tree.delete(item)

                # 更新分页按钮状态
                if results:
                    self.next_btn.config(state='normal')
                    self.prev_btn.config(state='normal' if page > 1 else 'disabled')
                else:
                    self.next_btn.config(state='disabled')
                    # Keep prev enabled if we are deep in pages? No, usually if no results, stop.
                    
                if not results:
                    self.search_status_var.set("未找到结果")
                    messagebox.showinfo("提示", f"在 {source} 中未找到关键词 '{query}' 的结果")
                    return
                    
                # 按分类分组
                categories = {}
                for res in results:
                    cat = res.get('category', 'Uncategorized')
                    if cat not in categories:
                        categories[cat] = []
                    categories[cat].append(res)
                
                for i, (cat, cat_books) in enumerate(categories.items()):
                    # 插入分类父节点
                    cat_id = f"search_cat_{i}"
                    self.search_tree.insert("", tk.END, iid=cat_id, values=("📂 " + cat, "", "", "", "", source, "分类"), open=True)
                    
                    for j, res in enumerate(cat_books):
                        self.search_tree.insert(cat_id, tk.END, iid=f"search_book_{i}_{j}", values=(
                            res.get('title', '未知'),
                            res.get('author', '未知'),
                            res.get('language', ''),
                            res.get('extension', ''),
                            res.get('size', ''),
                            res.get('source', source),
                            cat
                        ))
                
                self.search_status_var.set(f"搜索完成，找到 {len(results)} 个结果 (第 {page} 页)")
                
            self.root.after(0, update_ui)
        except Exception as e:
            self.root.after(0, lambda: messagebox.showerror("搜索错误", str(e)))
            self.root.after(0, lambda: self.search_status_var.set("搜索出错"))

    def on_search_result_select(self, event):
        """选择搜索结果"""
        selection = self.search_tree.selection()
        if not selection:
            return
            
        # 如果选中的是分类文件夹，不执行后续逻辑
        item_id = selection[0]
        if "_cat_" in item_id:
            self.selected_result = None
            self.search_detail_var.set("请选择具体的书籍")
            self.download_btn.config(state="disabled")
            return

        # 需要根据 Treeview 的组织方式找到对应的结果对象
        # 简单方案：在结果插入时将 index 存在 values 里，或者根据标题匹配
        # 这里我们遍历缓存寻找匹配的标题（最简单）
        item_values = self.search_tree.item(item_id, "values")
        title = item_values[0]
        
        self.selected_result = None
        for res in self.current_search_results:
            if res.get('title') == title:
                self.selected_result = res
                break
        
        if self.selected_result:
            res = self.selected_result
            detail = (
                f"标题: {res.get('title')}\n"
                f"作者: {res.get('author')}\n"
                f"分类: {res.get('category', 'Uncategorized')}\n"
                f"格式: {res.get('extension')} | 大小: {res.get('size')} | 语言: {res.get('language')}\n"
                f"来源: {res.get('source')}\n"
                f"详情: {res.get('metadata', '')}"
            )
            self.search_detail_var.set(detail)
            self.download_btn.config(state="normal")

    def on_download_click(self):
        """点击下载按钮"""
        if not self.selected_result:
            return
            
        res = self.selected_result
        self.search_status_var.set(f"正在下载: {res.get('title')}...")
        self.download_btn.config(state="disabled")
        
        # 开启线程下载
        threading.Thread(target=self.perform_download, args=(res,), daemon=True).start()

    def perform_download(self, result_item):
        """执行下载逻辑（后台线程）"""
        # 导入自定义异常
        from online_search import CloudflareError
        import webbrowser

        try:
            def progress_cb(current, total):
                if total > 0:
                    percent = (current / total) * 100
                    self.root.after(0, lambda: self.search_status_var.set(f"下载中: {percent:.1f}%"))
                else:
                    self.root.after(0, lambda: self.search_status_var.set(f"下载中: {current/1024/1024:.1f} MB"))
            
            file_path = self.online_search_manager.download_book(result_item, progress_callback=progress_cb)
            
            if file_path and os.path.exists(file_path):
                self.root.after(0, lambda: self.search_status_var.set("下载成功，正在导入..."))
                
                def load_downloaded():
                    self.file_path_var.set(file_path)
                    self.load_file_content(file_path)
                    # 切换到原文标签页
                    for i in range(self.notebook.index("end")):
                        if self.notebook.tab(i, "text") == "原文":
                            self.notebook.select(i)
                            break
                    messagebox.showinfo("成功", f"书籍已下载并成功导入：\n{os.path.basename(file_path)}")
                
                self.root.after(0, load_downloaded)
            else:
                self.root.after(0, lambda: messagebox.showerror("下载失败", "无法完成下载，文件未保存。"))
                self.root.after(0, lambda: self.search_status_var.set("下载失败"))

        except CloudflareError as ce:
            # 捕获 Cloudflare 错误，提示用户使用浏览器打开
            msg = "自动下载被 Cloudflare 拦截 (403 Forbidden)。\n这通常是因为网站启用了反爬虫保护。\n\n是否在浏览器中打开下载页面？"
            self.root.after(0, lambda: self._prompt_browser_open(msg, str(ce)))
            self.root.after(0, lambda: self.search_status_var.set("需要浏览器下载"))

        except Exception as e:
            self.root.after(0, lambda: messagebox.showerror("下载错误", str(e)))
            self.root.after(0, lambda: self.search_status_var.set("下载出错"))
        finally:
            self.root.after(0, lambda: self.download_btn.config(state="normal"))

    def _prompt_browser_open(self, msg, url):
        """提示并在浏览器打开链接"""
        import webbrowser
        if messagebox.askyesno("下载受阻", msg):
            webbrowser.open(url)

    def open_online_config(self):
        """打开在线搜索配置对话框"""
        config_window = tk.Toplevel(self.root)
        config_window.title("在线搜索配置")
        config_window.geometry("550x450")
        config_window.transient(self.root)
        config_window.grab_set()

        frame = ttk.Frame(config_window, padding="20")
        frame.pack(fill=tk.BOTH, expand=True)
        
        # Z-Library 配置
        zlib_frame = ttk.LabelFrame(frame, text="Z-Library 配置", padding="10")
        zlib_frame.pack(fill=tk.X, pady=(0, 10))
        
        zlib_config = self.config_manager.get('online_search.zlibrary', {})
        
        ttk.Label(zlib_frame, text="邮箱:").grid(row=0, column=0, sticky=tk.W, pady=2)
        email_var = tk.StringVar(value=zlib_config.get('email', ''))
        ttk.Entry(zlib_frame, textvariable=email_var, width=30).grid(row=0, column=1, sticky=tk.W, pady=2, padx=5)
        
        ttk.Label(zlib_frame, text="密码:").grid(row=1, column=0, sticky=tk.W, pady=2)
        pass_var = tk.StringVar(value=zlib_config.get('password', ''))
        ttk.Entry(zlib_frame, textvariable=pass_var, show="*", width=30).grid(row=1, column=1, sticky=tk.W, pady=2, padx=5)
        
        ttk.Label(zlib_frame, text="域名:").grid(row=2, column=0, sticky=tk.W, pady=2)
        domain_var = tk.StringVar(value=zlib_config.get('domain', 'https://singlelogin.re'))
        domain_entry = ttk.Entry(zlib_frame, textvariable=domain_var, width=30)
        domain_entry.grid(row=2, column=1, sticky=tk.W, pady=2, padx=5)

        def test_zlib_connection():
            """测试 Z-Library 连接"""
            email = email_var.get().strip()
            password = pass_var.get().strip()
            domain = domain_var.get().strip()
            
            if not email or not password:
                messagebox.showwarning("提示", "请输入邮箱和密码以测试登录")
                return

            self.config_manager.set('online_search.zlibrary.email', email, save=False)
            self.config_manager.set('online_search.zlibrary.password', password, save=False)
            self.config_manager.set('online_search.zlibrary.domain', domain, save=False)
            
            def run_test():
                try:
                    success = self.online_search_manager.login_zlibrary()
                    if success:
                        self.root.after(0, lambda: messagebox.showinfo("成功", "Z-Library 登录成功！"))
                    else:
                        self.root.after(0, lambda: messagebox.showerror("失败", "Z-Library 登录失败，请检查账号密码或域名"))
                except Exception as e:
                    self.root.after(0, lambda: messagebox.showerror("错误", f"连接出错: {str(e)}"))
            
            threading.Thread(target=run_test, daemon=True).start()

        ttk.Button(zlib_frame, text="测试登录", command=test_zlib_connection).grid(row=3, column=1, sticky=tk.E, pady=5, padx=5)
        
        # Anna's Archive 配置
        annas_frame = ttk.LabelFrame(frame, text="Anna's Archive 配置", padding="10")
        annas_frame.pack(fill=tk.X, pady=(0, 10))
        
        annas_config = self.config_manager.get('online_search.annas_archive', {})
        ttk.Label(annas_frame, text="域名:").grid(row=0, column=0, sticky=tk.W, pady=2)
        annas_domain_var = tk.StringVar(value=annas_config.get('domain', 'https://annas-archive.li'))
        annas_domain_entry = ttk.Entry(annas_frame, textvariable=annas_domain_var, width=30)
        annas_domain_entry.grid(row=0, column=1, sticky=tk.W, pady=2, padx=5)

        def test_annas_connection():
            """测试 Anna's Archive 连接"""
            domain = annas_domain_var.get().strip()
            
            def run_test():
                try:
                    import requests
                    resp = requests.get(domain, timeout=10)
                    if resp.status_code == 200:
                        self.root.after(0, lambda: messagebox.showinfo("成功", f"成功连接到 Anna's Archive！\n状态码: {resp.status_code}"))
                    else:
                        self.root.after(0, lambda: messagebox.showwarning("警告", f"服务器返回状态码: {resp.status_code}"))
                except Exception as e:
                    self.root.after(0, lambda: messagebox.showerror("错误", f"连接失败: {str(e)}"))

            threading.Thread(target=run_test, daemon=True).start()

        ttk.Button(annas_frame, text="测试连接", command=test_annas_connection).grid(row=1, column=1, sticky=tk.E, pady=5, padx=5)

        # 自动检测镜像功能
        def auto_detect_mirrors():
            """自动检测最快镜像"""
            btn = detect_btn
            btn.config(state='disabled', text="检测中...")
            
            def run_detection():
                try:
                    results = self.online_search_manager.check_mirrors()
                    
                    # 查找最佳镜像
                    best_annas = None
                    for m in results['annas_archive']:
                        if m['status'] == 'OK':
                            best_annas = m
                            break
                            
                    best_zlib = None
                    for m in results['zlibrary']:
                        if m['status'] == 'OK':
                            best_zlib = m
                            break
                    
                    msg = "检测结果:\n\n"
                    
                    if best_annas:
                        msg += f"Anna's Archive (最快): {best_annas['url']} ({best_annas['latency']}ms)\n"
                        self.root.after(0, lambda: annas_domain_var.set(best_annas['url']))
                    else:
                        msg += "Anna's Archive: 未找到可用镜像\n"
                        
                    if best_zlib:
                        msg += f"Z-Library (最快): {best_zlib['url']} ({best_zlib['latency']}ms)\n"
                        self.root.after(0, lambda: domain_var.set(best_zlib['url']))
                    else:
                        msg += "Z-Library: 未找到可用镜像\n"
                        
                    self.root.after(0, lambda: messagebox.showinfo("镜像检测完成", msg))
                    
                except Exception as e:
                    self.root.after(0, lambda: messagebox.showerror("错误", f"检测失败: {e}"))
                finally:
                    self.root.after(0, lambda: btn.config(state='normal', text="自动检测镜像"))

            threading.Thread(target=run_detection, daemon=True).start()

        detect_btn = ttk.Button(frame, text="自动检测并选择最快镜像", command=auto_detect_mirrors)
        detect_btn.pack(side=tk.LEFT, padx=10, pady=10)

        def save_online_config():
            self.config_manager.set('online_search.zlibrary.email', email_var.get().strip())
            self.config_manager.set('online_search.zlibrary.password', pass_var.get().strip())
            self.config_manager.set('online_search.zlibrary.domain', domain_var.get().strip())
            self.config_manager.set('online_search.annas_archive.domain', annas_domain_var.get().strip())
            messagebox.showinfo("成功", "在线搜索配置已保存")
            config_window.destroy()

        ttk.Button(frame, text="保存配置", command=save_online_config).pack(side=tk.RIGHT, pady=10)

    def create_menu_bar(self):
        """创建顶部菜单栏"""
        menubar = tk.Menu(self.root)
        self.root.config(menu=menubar)

        # File Menu
        file_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="文件 (File)", menu=file_menu)
        file_menu.add_command(label="打开文件...", command=self.browse_file)
        file_menu.add_command(label="从 URL 导入网页...", command=self.import_from_url)
        file_menu.add_command(label="从剪贴板导入文本", command=self.import_from_clipboard)
        file_menu.add_separator()
        file_menu.add_command(label="退出", command=self.on_closing)

        # Tools Menu
        tools_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="工具 (Tools)", menu=tools_menu)
        tools_menu.add_command(label="智能提取术语 (Auto Glossary)", command=self.generate_glossary_action)
        tools_menu.add_command(label="翻译记忆库编辑器 (TM Editor)", command=self.open_tm_editor)
        tools_menu.add_command(label="格式转换工具箱 (Format Converter)", command=self.open_format_converter)
        tools_menu.add_command(label="云端分享 (Upload & Share)", command=self.open_cloud_share)
        tools_menu.add_separator()
        tools_menu.add_command(label="导出双语对照 Word (.docx)", command=self.export_bilingual_docx_action)
        tools_menu.add_command(label="生成有声书 (Audiobook)", command=self.export_audiobook)
        
        # View Menu
        view_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="视图 (View)", menu=view_menu)
        view_menu.add_command(label="切换主题 (明亮/暗黑)", command=self.toggle_theme)

    def import_from_url(self):
        """从 URL 导入网页内容"""
        url = simpledialog.askstring("导入网页", "请输入网页 URL:")
        if not url: return
        
        try:
            self.progress_text_var.set("正在抓取网页...")
            # Run in thread to avoid freezing
            def fetch_thread():
                try:
                    title, content = self.web_importer.fetch_content(url)
                    self.root.after(0, lambda: self._load_imported_content(f"URL: {title}", content))
                except Exception as e:
                    self.root.after(0, lambda: messagebox.showerror("导入失败", str(e)))
                    self.root.after(0, lambda: self.progress_text_var.set("导入失败"))
            
            threading.Thread(target=fetch_thread, daemon=True).start()
            
        except Exception as e:
            messagebox.showerror("导入失败", str(e))

    def import_from_clipboard(self):
        """从剪贴板导入文本"""
        try:
            content = self.root.clipboard_get()
            if not content.strip():
                messagebox.showwarning("提示", "剪贴板为空")
                return
            self._load_imported_content("Clipboard Content", content)
            
        except Exception as e:
            messagebox.showerror("错误", f"无法读取剪贴板: {e}")

    def _load_imported_content(self, title, content):
        """Helper to load content into the editor"""
        self.load_content_into_workspace(
            title=title,
            content=content,
            filepath=None,
            clear_progress_cache=True,
        )
        self.file_info_var.set(f"已加载: {title[:20]}... ({len(content)} 字符)")
        self.progress_text_var.set("导入成功")

    def generate_glossary_action(self):
        """智能提取术语"""
        if not self.current_text:
            messagebox.showwarning("警告", "请先加载文本")
            return
            
        if not messagebox.askyesno("确认", "这将使用 LLM 分析文本前 4000 字并提取术语，可能消耗少量 Token。\n是否继续？"):
            return
            
        def run_extraction():
            try:
                self.root.after(0, lambda: self.progress_text_var.set("正在分析文本提取术语..."))
                terms = self.smart_glossary.extract_terms(self.current_text)
                
                if not terms:
                    self.root.after(0, lambda: messagebox.showinfo("结果", "未提取到重要术语"))
                    return
                    
                self.root.after(0, lambda: self._show_glossary_import_dialog(terms))
                
            except Exception as e:
                self.root.after(0, lambda: messagebox.showerror("错误", str(e)))
            finally:
                self.root.after(0, lambda: self.progress_text_var.set("就绪"))
                
        threading.Thread(target=run_extraction, daemon=True).start()

    def _show_glossary_import_dialog(self, terms):
        """显示术语导入确认对话框"""
        win = tk.Toplevel(self.root)
        win.title("提取到的术语")
        win.geometry("600x400")
        
        ttk.Label(win, text="勾选要添加到当前术语表的词条:").pack(pady=5)
        
        frame = ttk.Frame(win)
        frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        canvas = tk.Canvas(frame)
        scrollbar = ttk.Scrollbar(frame, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)
        
        scrollable_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        check_vars = []
        for i, (term, trans, type_) in enumerate(terms):
            var = tk.BooleanVar(value=True)
            check_vars.append((var, term, trans, type_))
            cb = ttk.Checkbutton(scrollable_frame, text=f"[{type_}] {term} -> {trans}", variable=var)
            cb.pack(anchor="w", padx=5, pady=2)
            
        def do_import():
            count = 0
            target_glossary = "Auto_Extracted"
            if not self.glossary_manager.load_glossary(target_glossary):
                self.glossary_manager.create_glossary(target_glossary, "AI 自动提取的术语")
            
            for var, term, trans, type_ in check_vars:
                if var.get():
                    self.glossary_manager.add_term(target_glossary, term, trans, notes=f"Type: {type_}")
                    count += 1
            
            messagebox.showinfo("成功", f"已导入 {count} 个术语到 '{target_glossary}' 表")
            win.destroy()
            
        ttk.Button(win, text="导入选中项", command=do_import).pack(pady=10)

    def export_bilingual_docx_action(self):
        """导出双语对照 Word 文档"""
        if not self.translated_segments:
            messagebox.showwarning("警告", "没有可导出的译文")
            return
            
        try:
            # Check for docx library
            try:
                import docx
            except ImportError:
                messagebox.showerror("错误", "未安装 python-docx 库")
                return

            filename = filedialog.asksaveasfilename(
                title="导出双语对照 Word",
                defaultextension=".docx",
                filetypes=[("Word 文档", "*.docx")]
            )
            if not filename: return
            
            # Use DocxHandler to create it (we instantiate a dummy one or use static method if we had one, 
            # but since we added it as instance method, we can create a temporary handler or just use the logic directly here)
            # Actually, reusing the logic I put in DocxHandler is best, but I need an instance.
            # I will just re-implement the simple logic here to avoid dependency on an existing file for DocxHandler init.
            
            from docx import Document
            from docx.shared import Pt
            
            doc = Document()
            doc.add_heading('双语对照翻译 / Bilingual Translation', 0)
            
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '原文 (Original)'
            hdr_cells[1].text = '译文 (Translation)'
            
            limit = max(len(self.source_segments), len(self.translated_segments))
            for i in range(limit):
                row_cells = table.add_row().cells
                
                # Source
                if i < len(self.source_segments):
                    row_cells[0].text = self.source_segments[i]
                
                # Target
                if i < len(self.translated_segments):
                    row_cells[1].text = self.translated_segments[i]
                    
            doc.save(filename)
            messagebox.showinfo("成功", f"双语文档已导出: {filename}")
            
        except Exception as e:
            messagebox.showerror("导出失败", str(e))

    def toggle_theme(self):
        """切换明亮/暗黑主题"""
        style = ttk.Style()
        
        if self.current_theme == "light":
            self.current_theme = "dark"
            style.theme_use('clam')
            
            bg_color = "#2b2b2b"
            fg_color = "#ffffff"
            field_bg = "#3c3f41"
            select_bg = "#005fb8"
            
            style.configure(".", background=bg_color, foreground=fg_color)
            style.configure("TLabel", background=bg_color, foreground=fg_color)
            style.configure("TButton", background=field_bg, foreground=fg_color, borderwidth=1)
            style.map("TButton", background=[("active", "#4c5052")])
            style.configure("TEntry", fieldbackground=field_bg, foreground=fg_color)
            style.configure("TCombobox", fieldbackground=field_bg, foreground=fg_color, background=bg_color)
            style.configure("Treeview", background=field_bg, foreground=fg_color, fieldbackground=field_bg)
            style.map("Treeview", background=[("selected", select_bg)])
            style.configure("TLabelframe", background=bg_color, foreground=fg_color)
            style.configure("TLabelframe.Label", background=bg_color, foreground=fg_color)
            
            # Text widgets
            for widget in [self.original_text, self.translated_text_widget, self.failed_source_text, 
                          self.manual_translation_text, self.analysis_text, self.comp_source_text, self.comp_target_text]:
                try:
                    widget.config(bg=field_bg, fg=fg_color, insertbackground=fg_color)
                except: pass
                
            # Listboxes
            for lb in [self.failed_listbox, self.analysis_listbox]:
                try:
                    lb.config(bg=field_bg, fg=fg_color)
                except: pass
            
        else:
            self.current_theme = "light"
            style.theme_use('default')
            
            # Reset Text widgets
            for widget in [self.original_text, self.translated_text_widget, self.failed_source_text, 
                          self.manual_translation_text, self.analysis_text, self.comp_source_text, self.comp_target_text]:
                try:
                    widget.config(bg="white", fg="black", insertbackground="black")
                except: pass
                
            # Reset Listboxes
            for lb in [self.failed_listbox, self.analysis_listbox]:
                try:
                    lb.config(bg="white", fg="black")
                except: pass

    def open_tm_editor(self):
        """打开翻译记忆库编辑器"""
        TMEditorDialog(self.root, self.translation_memory)

    def open_format_converter(self):
        """打开格式转换工具箱"""
        def load_callback(path):
            if path and os.path.exists(path):
                self.root.after(0, lambda: self.file_path_var.set(path))
                self.root.after(0, lambda: self.load_file_content(path))
                
        FormatConverterDialog(self.root, load_callback)

def main():
    """主程序入口"""
    # 设置日志记录
    class Logger(object):
        def __init__(self, filename="translator.log"):
            self.terminal = sys.stdout
            self.log = open(filename, "a", encoding="utf-8")

        def write(self, message):
            # 避免写入空行或只有换行符的行（可选）
            self.terminal.write(message)
            self.log.write(message)
            self.log.flush()

        def flush(self):
            self.terminal.flush()
            self.log.flush()

    # 重定向输出到日志文件
    sys.stdout = Logger()
    sys.stderr = Logger("translator_error.log")
    
    print(f"\n{'='*50}")
    print(f"启动时间: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"版本: {CONFIG_VERSION}")
    print(f"{'='*50}\n")

    # 检查依赖
    missing_libs = []
    if not PDF_SUPPORT:
        missing_libs.append("PyPDF2 (用于PDF支持)")
    if not EPUB_SUPPORT:
        missing_libs.append("ebooklib, beautifulsoup4 (用于EPUB支持)")
    if not GEMINI_SUPPORT:
        missing_libs.append("google-generativeai (用于Gemini API)")
    if not OPENAI_SUPPORT:
        missing_libs.append("openai (用于OpenAI API)")
    if not REQUESTS_SUPPORT:
        missing_libs.append("requests (用于自定义API)")

    if missing_libs:
        print("=" * 60)
        print("警告: 以下库未安装，部分功能将不可用:")
        for lib in missing_libs:
            print(f"  - {lib}")
        print("\n安装命令:")
        print("py -m pip install PyPDF2 ebooklib beautifulsoup4 google-generativeai openai requests")
        print("=" * 60)
        print()

    root = tk.Tk()
    app = BookTranslatorGUI(root)
    root.mainloop()


if __name__ == "__main__":
    main()
