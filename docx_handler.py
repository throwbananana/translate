#! python
# -*- coding: utf-8 -*-
"""
Docx 处理模块 (Docx Handler)
用于读取 Word 文档并在翻译后尽量保留原格式回填。

说明：
- 该实现优先保证“可稳定导出”，而不是复杂 run 级精确保真。
- 段落与表格单元格段落都支持基础回填。
- 输出目录会自动创建，减少导出阶段的路径错误。
"""

from pathlib import Path

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocxHandler:
    def __init__(self, filepath):
        if not DOCX_AVAILABLE:
            raise ImportError("需安装 python-docx 库才能使用此功能: py -m pip install python-docx")

        self.filepath = filepath
        self.document = Document(filepath)
        self.para_map = []
        self._scan_document()

    def _scan_document(self):
        """扫描文档，建立顺序映射。"""
        self.para_map = []

        for i, para in enumerate(self.document.paragraphs):
            if para.text.strip():
                self.para_map.append({
                    'type': 'paragraph',
                    'index': i,
                    'original_text': para.text,
                })

        for table_idx, table in enumerate(self.document.tables):
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    for para_idx, para in enumerate(cell.paragraphs):
                        if para.text.strip():
                            self.para_map.append({
                                'type': 'table',
                                'table_idx': table_idx,
                                'row_idx': row_idx,
                                'col_idx': col_idx,
                                'cell_p_idx': para_idx,
                                'original_text': para.text,
                            })

    def extract_text(self):
        """提取纯文本用于翻译。"""
        return "\n\n".join(item['original_text'] for item in self.para_map)

    def _resolve_target_paragraph(self, doc, mapping):
        if mapping['type'] == 'paragraph':
            return doc.paragraphs[mapping['index']]

        if mapping['type'] == 'table':
            return (
                doc.tables[mapping['table_idx']]
                .rows[mapping['row_idx']]
                .cells[mapping['col_idx']]
                .paragraphs[mapping['cell_p_idx']]
            )

        raise ValueError(f"未知段落类型: {mapping['type']}")

    def _normalize_segments(self, segments, argument_name):
        """将外部输入标准化为字符串列表，减少导出阶段的类型错误。"""
        if isinstance(segments, str):
            raise TypeError(f"{argument_name} 必须是段落列表，不能直接传入单个字符串")

        if segments is None:
            raise TypeError(f"{argument_name} 不能为空")

        try:
            normalized = ["" if item is None else str(item) for item in segments]
        except TypeError as exc:
            raise TypeError(f"{argument_name} 必须是可迭代的段落列表") from exc

        return normalized

    def _ensure_output_dir(self, output_path):
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        return output_file

    def _replace_paragraph_text(self, paragraph, text):
        """尽量保留基础段落样式并替换文本。"""
        runs = list(paragraph.runs)
        if not runs:
            paragraph.add_run(text)
            return

        runs[0].text = text
        for run in runs[1:]:
            run.text = ''

    def save_translated_file(self, translated_segments, output_path):
        """将翻译后的段落回填到文档中并保存。"""
        translated_segments = self._normalize_segments(translated_segments, 'translated_segments')

        if not translated_segments:
            raise ValueError("没有可写入的译文段落")

        if len(translated_segments) != len(self.para_map):
            print(
                f"警告: 翻译段落数 ({len(translated_segments)}) 与原文有效段落数 ({len(self.para_map)}) 不匹配，"
                "将按最小长度回填。"
            )

        doc = Document(self.filepath)
        limit = min(len(translated_segments), len(self.para_map))

        for idx in range(limit):
            mapping = self.para_map[idx]
            trans_text = translated_segments[idx]
            try:
                paragraph = self._resolve_target_paragraph(doc, mapping)
                self._replace_paragraph_text(paragraph, trans_text)
            except Exception as exc:
                print(f"DOCX 回填失败 (段落 {idx + 1}): {exc}")

        output_file = self._ensure_output_dir(output_path)
        doc.save(str(output_file))
        return str(output_file)

    def save_bilingual_file(self, original_segments, translated_segments, output_path):
        """保存为双语对照 Word 文档。"""
        original_segments = self._normalize_segments(original_segments, 'original_segments')
        translated_segments = self._normalize_segments(translated_segments, 'translated_segments')

        doc = Document()
        doc.add_heading('双语对照翻译 / Bilingual Translation', 0)

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '原文 (Original)'
        hdr_cells[1].text = '译文 (Translation)'

        limit = max(len(original_segments), len(translated_segments))
        for i in range(limit):
            row_cells = table.add_row().cells
            if i < len(original_segments):
                row_cells[0].text = original_segments[i]
            if i < len(translated_segments):
                row_cells[1].text = translated_segments[i]

        output_file = self._ensure_output_dir(output_path)
        doc.save(str(output_file))
        return str(output_file)
