#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文件处理器 (File Processor) 模块
支持多种文件格式的读取、文本提取和智能分段

支持的格式：
- TXT: 自动编码检测
- PDF: PyPDF2 文本提取
- EPUB: ebooklib + BeautifulSoup + OCR
- DOCX: python-docx (新增)
- Markdown: 原生支持 (新增)
- RTF: 基础支持 (新增)
"""

import os
import re
from pathlib import Path
from typing import List, Optional, Callable, Tuple

# PDF 支持
try:
    import PyPDF2
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    import pdfplumber
    PDFPLUMBER_SUPPORT = True
except ImportError:
    PDFPLUMBER_SUPPORT = False

try:
    from pdf2image import convert_from_path
    PDF2IMAGE_SUPPORT = True
except ImportError:
    PDF2IMAGE_SUPPORT = False

# EPUB 支持
try:
    import ebooklib
    from ebooklib import epub
    from bs4 import BeautifulSoup
    EPUB_SUPPORT = True
except ImportError:
    EPUB_SUPPORT = False

# DOCX 支持
try:
    from docx import Document as DocxDocument
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

# OCR 支持
try:
    import pytesseract
    from PIL import Image
    import io
    OCR_SUPPORT = True
except ImportError:
    OCR_SUPPORT = False


class FileProcessor:
    """文件处理器"""

    # 默认分段大小
    DEFAULT_SEGMENT_SIZE = 800

    # 支持的文件格式
    SUPPORTED_FORMATS = {
        '.txt': 'TXT',
        '.md': 'Markdown',
        '.markdown': 'Markdown',
        '.pdf': 'PDF',
        '.epub': 'EPUB',
        '.docx': 'DOCX',
        '.rtf': 'RTF'
    }

    def __init__(self, segment_size: int = None):
        """
        初始化文件处理器

        Args:
            segment_size: 默认分段大小
        """
        self.segment_size = segment_size or self.DEFAULT_SEGMENT_SIZE

        # 尝试自动配置 Tesseract 路径 (Windows)
        if OCR_SUPPORT and os.name == 'nt':
            self._auto_configure_tesseract()

    def _auto_configure_tesseract(self):
        """自动配置 Tesseract 路径"""
        try:
            import subprocess
            subprocess.run(['tesseract', '--version'],
                         stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True)
            return
        except (FileNotFoundError, Exception):
            pass

        # 常见安装路径
        common_paths = [
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
            os.path.expandvars(r"%LOCALAPPDATA%\Tesseract-OCR\tesseract.exe"),
            os.path.expandvars(r"%PROGRAMFILES%\Tesseract-OCR\tesseract.exe")
        ]

        for path in common_paths:
            if os.path.exists(path):
                print(f"Auto-detected Tesseract at: {path}")
                pytesseract.pytesseract.tesseract_cmd = path
                return

    @classmethod
    def get_supported_formats(cls) -> dict:
        """获取当前支持的文件格式"""
        formats = {'.txt': True, '.md': True, '.markdown': True}

        if PDF_SUPPORT or PDFPLUMBER_SUPPORT:
            formats['.pdf'] = True
        if EPUB_SUPPORT:
            formats['.epub'] = True
        if DOCX_SUPPORT:
            formats['.docx'] = True

        # RTF 基础支持总是可用
        formats['.rtf'] = True

        return formats

    @classmethod
    def get_file_filter(cls) -> List[Tuple[str, str]]:
        """获取用于文件对话框的过滤器"""
        supported = cls.get_supported_formats()
        extensions = ' '.join(f'*{ext}' for ext in supported.keys())

        filters = [("所有支持的文件", extensions)]

        if '.pdf' in supported:
            filters.append(("PDF文件", "*.pdf"))
        if '.epub' in supported:
            filters.append(("EPUB文件", "*.epub"))
        if '.docx' in supported:
            filters.append(("Word文档", "*.docx"))
        filters.append(("文本文件", "*.txt"))
        filters.append(("Markdown文件", "*.md *.markdown"))

        return filters

    def read_file(self, filepath: str,
                  progress_callback: Callable[[str], None] = None) -> str:
        """
        读取文件内容

        Args:
            filepath: 文件路径
            progress_callback: 进度回调函数

        Returns:
            文件内容字符串
        """
        ext = os.path.splitext(filepath)[1].lower()

        if ext == '.pdf':
            if not PDF_SUPPORT and not PDFPLUMBER_SUPPORT:
                raise ImportError("未安装 PyPDF2 或 pdfplumber，无法读取 PDF 文件。")
            return self.extract_pdf_text(filepath, progress_callback)

        elif ext == '.epub':
            if not EPUB_SUPPORT:
                raise ImportError("未安装 ebooklib 或 beautifulsoup4，无法读取 EPUB 文件。请运行: pip install ebooklib beautifulsoup4")
            return self.extract_epub_text(filepath, progress_callback)

        elif ext == '.docx':
            if not DOCX_SUPPORT:
                raise ImportError("未安装 python-docx，无法读取 DOCX 文件。请运行: pip install python-docx")
            return self.extract_docx_text(filepath, progress_callback)

        elif ext in ['.md', '.markdown']:
            return self.read_markdown_file(filepath, progress_callback)

        elif ext == '.rtf':
            return self.read_rtf_file(filepath, progress_callback)

        else:
            # 默认当做文本文件处理
            return self.read_text_file(filepath, progress_callback)

    def read_text_file(self, filepath: str,
                       progress_callback: Callable[[str], None] = None) -> str:
        """
        读取文本文件（自动检测编码）

        Args:
            filepath: 文件路径
            progress_callback: 进度回调

        Returns:
            文件内容
        """
        if progress_callback:
            progress_callback("正在检测文件编码...")

        encodings = ['utf-8', 'utf-8-sig', 'gbk', 'gb18030', 'utf-16', 'latin1', 'cp1252']

        for encoding in encodings:
            try:
                with open(filepath, 'r', encoding=encoding) as f:
                    content = f.read()
                if progress_callback:
                    progress_callback(f"使用 {encoding} 编码读取成功")
                return content
            except UnicodeDecodeError:
                continue

        raise ValueError("无法识别文件编码，请确保文件为有效的文本文件")

    def read_markdown_file(self, filepath: str,
                           progress_callback: Callable[[str], None] = None) -> str:
        """
        读取 Markdown 文件

        Args:
            filepath: 文件路径
            progress_callback: 进度回调

        Returns:
            文件内容（保留 Markdown 格式）
        """
        if progress_callback:
            progress_callback("正在读取 Markdown 文件...")

        content = self.read_text_file(filepath)

        if progress_callback:
            progress_callback("Markdown 文件读取完成")

        return content

    def read_rtf_file(self, filepath: str,
                      progress_callback: Callable[[str], None] = None) -> str:
        """
        读取 RTF 文件（基础支持）

        Args:
            filepath: 文件路径
            progress_callback: 进度回调

        Returns:
            提取的文本内容
        """
        if progress_callback:
            progress_callback("正在读取 RTF 文件...")

        with open(filepath, 'rb') as f:
            content = f.read()

        # 尝试解码
        try:
            text = content.decode('utf-8')
        except:
            try:
                text = content.decode('latin1')
            except:
                text = content.decode('cp1252', errors='ignore')

        # 简单的 RTF 文本提取
        # 移除 RTF 控制字符
        text = re.sub(r'\\[a-z]+\d*\s?', '', text)
        text = re.sub(r'\{|\}', '', text)
        text = re.sub(r'\\\'[0-9a-f]{2}', '', text)

        # 清理多余空白
        text = re.sub(r'\s+', ' ', text)
        text = text.strip()

        if progress_callback:
            progress_callback("RTF 文件读取完成")

        return text

    def extract_pdf_text(self, filepath: str,
                         progress_callback: Callable[[str], None] = None) -> str:
        """
        提取 PDF 文本 (增强版：支持 pdfplumber 布局保持和 OCR 回退)

        Args:
            filepath: PDF 文件路径
            progress_callback: 进度回调

        Returns:
            提取的文本
        """
        text_parts = []
        used_method = "Unknown"

        # 优先使用 pdfplumber (更好的排版处理)
        if PDFPLUMBER_SUPPORT:
            try:
                with pdfplumber.open(filepath) as pdf:
                    total_pages = len(pdf.pages)
                    used_method = "pdfplumber"
                    
                    for i, page in enumerate(pdf.pages):
                        if progress_callback:
                            progress_callback(f"正在读取PDF (pdfplumber)... {i+1}/{total_pages}")
                        
                        page_text = page.extract_text()
                        
                        # 检测扫描件：如果文字很少，尝试 OCR
                        if (not page_text or len(page_text.strip()) < 50) and OCR_SUPPORT:
                            ocr_text = self._ocr_pdf_page(filepath, i+1)
                            if ocr_text:
                                page_text = f"[OCR识别-P{i+1}]\n{ocr_text}"
                                if progress_callback:
                                    progress_callback(f"已对第 {i+1} 页进行 OCR 识别")

                        if page_text:
                            text_parts.append(page_text)
                            
                return '\n\n'.join(text_parts)
            except Exception as e:
                print(f"pdfplumber 读取失败: {e}，将回退到 PyPDF2")
        
        # 回退到 PyPDF2
        if PDF_SUPPORT:
            try:
                with open(filepath, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    total_pages = len(pdf_reader.pages)
                    used_method = "PyPDF2"

                    for page_num, page in enumerate(pdf_reader.pages):
                        if progress_callback:
                            progress_callback(f"正在读取PDF (PyPDF2)... {page_num + 1}/{total_pages}")

                        page_text = page.extract_text()
                        
                        # PyPDF2 扫描件检测与 OCR 回退
                        if (not page_text or len(page_text.strip()) < 50) and OCR_SUPPORT:
                             # PyPDF2 不容易直接转图片，依然尝试用 pdf2image (如果可用)
                             ocr_text = self._ocr_pdf_page(filepath, page_num + 1)
                             if ocr_text:
                                 page_text = f"[OCR识别-P{page_num+1}]\n{ocr_text}"
                                 if progress_callback:
                                     progress_callback(f"已对第 {page_num+1} 页进行 OCR 识别")

                        if page_text:
                            text_parts.append(page_text)

                return '\n\n'.join(text_parts)
            except Exception as e:
                print(f"PyPDF2 读取失败: {e}")
                raise e
        else:
             raise ImportError("未安装 PDF 读取库 (pdfplumber 或 PyPDF2)")

    def _ocr_pdf_page(self, filepath, page_num, dpi=200):
        """辅助方法：对 PDF 指定页进行 OCR"""
        if not OCR_SUPPORT or not PDF2IMAGE_SUPPORT:
            return ""
        
        try:
            # pdf2image 的 page_num 是从 1 开始的吗？convert_from_path 的 first_page 是 1-based
            images = convert_from_path(
                filepath, 
                first_page=page_num, 
                last_page=page_num,
                dpi=dpi
            )
            if images:
                # 默认只取第一张（一页PDF通常转为一张图）
                return pytesseract.image_to_string(images[0], lang='chi_sim+eng').strip()
        except Exception as e:
            # 常见的错误是 Poppler 未安装
            # 只有第一次报错时打印，避免刷屏
            if "poppler" in str(e).lower() and page_num == 1:
                print(f"OCR 警告: 无法将 PDF 转为图片，可能是未安装 Poppler。\n详情: {e}")
            return ""
        return ""

    def extract_docx_text(self, filepath: str,
                          progress_callback: Callable[[str], None] = None) -> str:
        """
        提取 DOCX 文本

        Args:
            filepath: DOCX 文件路径
            progress_callback: 进度回调

        Returns:
            提取的文本
        """
        if progress_callback:
            progress_callback("正在读取 Word 文档...")

        doc = DocxDocument(filepath)
        text_parts = []

        total_paragraphs = len(doc.paragraphs)

        for idx, para in enumerate(doc.paragraphs):
            if para.text.strip():
                text_parts.append(para.text)

            if progress_callback and (idx + 1) % 50 == 0:
                progress_callback(f"正在读取DOCX... {idx + 1}/{total_paragraphs}")

        # 也提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(' | '.join(row_text))

        if progress_callback:
            progress_callback("Word 文档读取完成")

        return '\n\n'.join(text_parts)

    def extract_epub_text(self, filepath: str,
                          progress_callback: Callable[[str], None] = None) -> str:
        """
        提取 EPUB 文本（带 OCR 支持）

        Args:
            filepath: EPUB 文件路径
            progress_callback: 进度回调

        Returns:
            提取的文本
        """
        book = epub.read_epub(filepath)
        text_parts = []

        items = list(book.get_items_of_type(ebooklib.ITEM_DOCUMENT))
        total_items = len(items)

        # 建立图片映射
        image_items = {}
        for item in book.get_items_of_type(ebooklib.ITEM_IMAGE):
            image_items[item.get_name()] = item
            image_items[os.path.basename(item.get_name())] = item

        for idx, item in enumerate(items):
            if progress_callback:
                progress_callback(f"正在读取EPUB... {idx + 1}/{total_items}")

            soup = BeautifulSoup(item.get_content(), 'html.parser')

            # 处理图片
            self._process_epub_images(soup, image_items)

            text_parts.append(soup.get_text())

        return '\n\n'.join(text_parts)

    def _process_epub_images(self, soup, image_items: dict):
        """处理 EPUB 中的图片（OCR 或提取 alt 文本）"""

        def try_get_ocr_text(src_href):
            if not src_href or not OCR_SUPPORT:
                return ""

            img_item = image_items.get(src_href) or \
                       image_items.get(os.path.basename(src_href))

            if img_item:
                ocr_text = self.perform_ocr(img_item.get_content())
                if ocr_text:
                    return f"[OCR识别: {' '.join(ocr_text.split())}]"
            return ""

        # 处理 <img> 标签
        for img in soup.find_all('img'):
            alt_text = img.get('alt', '').strip()
            title_text = img.get('title', '').strip()
            aria_label = img.get('aria-label', '').strip()

            final_text = alt_text or title_text or aria_label

            if not final_text:
                src = img.get('src', '')
                final_text = try_get_ocr_text(src)

            if not final_text:
                src = img.get('src', '')
                if src:
                    final_text = f"[图片: {os.path.basename(src)}]"

            if final_text:
                img.replace_with(f"\n{final_text}\n")

        # 处理 SVG <image> 标签
        for svg_img in soup.find_all('image'):
            aria_label = svg_img.get('aria-label', '').strip()
            parent_title = ""

            if svg_img.parent and svg_img.parent.name == 'svg':
                title_tag = svg_img.parent.find('title')
                if title_tag:
                    parent_title = title_tag.get_text().strip()

            final_text = aria_label or parent_title

            if not final_text:
                href = svg_img.get('xlink:href') or svg_img.get('href')
                final_text = try_get_ocr_text(href)

            if not final_text:
                href = svg_img.get('xlink:href') or svg_img.get('href')
                if href:
                    final_text = f"[图片: {os.path.basename(href)}]"

            if final_text and svg_img.parent:
                new_tag = soup.new_tag("p")
                new_tag.string = final_text
                svg_img.parent.insert_after(new_tag)

    def perform_ocr(self, image_content: bytes, lang: str = 'chi_sim+eng') -> str:
        """
        对图片进行 OCR 识别

        Args:
            image_content: 图片二进制数据
            lang: 识别语言，默认简体中文+英文

        Returns:
            识别出的文本
        """
        if not OCR_SUPPORT:
            return ""

        try:
            image = Image.open(io.BytesIO(image_content))
            text = pytesseract.image_to_string(image, lang=lang)
            return text.strip()

        except Exception as e:
            error_str = str(e)
            print(f"OCR Error (lang={lang}): {error_str}")

            # 尝试回退到英语
            if "tessdata" in error_str or "language" in error_str.lower():
                if lang != 'eng':
                    try:
                        image = Image.open(io.BytesIO(image_content))
                        text = pytesseract.image_to_string(image, lang='eng')
                        return text.strip()
                    except Exception as e2:
                        print(f"OCR Fallback Error: {e2}")

            return ""

    def split_text_into_segments(self, text: str,
                                  max_length: int = None,
                                  preserve_paragraphs: bool = True) -> List[str]:
        """
        将文本分割成段落

        Args:
            text: 完整文本
            max_length: 每段最大长度（默认使用实例配置）
            preserve_paragraphs: 是否尽量保持段落完整

        Returns:
            段落列表
        """
        if max_length is None:
            max_length = self.segment_size

        if not text:
            return []

        segments = []

        if preserve_paragraphs:
            # 按段落分割
            paragraphs = re.split(r'\n\s*\n', text)
        else:
            paragraphs = text.split('\n')

        current_segment = ""

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            # 如果当前段落加上新段落不超过限制，合并
            if len(current_segment) + len(para) + 2 <= max_length:
                if current_segment:
                    current_segment += "\n\n" + para
                else:
                    current_segment = para
            else:
                # 保存当前段落
                if current_segment:
                    segments.append(current_segment.strip())

                # 如果新段落本身超长，需要强制分割
                if len(para) > max_length:
                    # 尝试按句子分割
                    sentences = self._split_into_sentences(para)
                    temp_segment = ""

                    for sentence in sentences:
                        if len(temp_segment) + len(sentence) + 1 <= max_length:
                            if temp_segment:
                                temp_segment += " " + sentence
                            else:
                                temp_segment = sentence
                        else:
                            if temp_segment:
                                segments.append(temp_segment.strip())

                            # 如果单个句子还是太长，强制按字符分割
                            if len(sentence) > max_length:
                                for i in range(0, len(sentence), max_length):
                                    segments.append(sentence[i:i+max_length])
                                temp_segment = ""
                            else:
                                temp_segment = sentence

                    if temp_segment:
                        current_segment = temp_segment
                    else:
                        current_segment = ""
                else:
                    current_segment = para

        # 添加最后一个段落
        if current_segment:
            segments.append(current_segment.strip())

        return segments

    def _split_into_sentences(self, text: str) -> List[str]:
        """将文本分割成句子"""
        # 支持中英文句号、问号、感叹号
        pattern = r'([。！？.!?]+)'
        parts = re.split(pattern, text)

        sentences = []
        current = ""

        for i, part in enumerate(parts):
            if re.match(pattern, part):
                current += part
                if current.strip():
                    sentences.append(current.strip())
                current = ""
            else:
                current = part

        if current.strip():
            sentences.append(current.strip())

        return sentences

    def detect_language(self, text: str) -> str:
        """
        简单的语言检测

        Args:
            text: 要检测的文本

        Returns:
            语言代码: 'zh', 'en', 'ja', 'mixed', 'unknown'
        """
        if not text or len(text.strip()) == 0:
            return 'unknown'

        # 统计各类字符
        chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
        japanese_chars = len(re.findall(r'[\u3040-\u30ff\u31f0-\u31ff]', text))
        latin_chars = len(re.findall(r'[A-Za-z]', text))
        total_chars = len(re.findall(r'\S', text)) or 1

        chinese_ratio = chinese_chars / total_chars
        japanese_ratio = japanese_chars / total_chars
        latin_ratio = latin_chars / total_chars

        if chinese_ratio > 0.5:
            return 'zh'
        elif japanese_ratio > 0.3:
            return 'ja'
        elif latin_ratio > 0.5:
            return 'en'
        elif chinese_ratio > 0.2 or japanese_ratio > 0.1 or latin_ratio > 0.2:
            return 'mixed'
        else:
            return 'unknown'

    def get_file_info(self, filepath: str) -> dict:
        """
        获取文件信息

        Args:
            filepath: 文件路径

        Returns:
            文件信息字典
        """
        path = Path(filepath)

        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {filepath}")

        stat = path.stat()
        ext = path.suffix.lower()

        return {
            'name': path.name,
            'path': str(path.absolute()),
            'extension': ext,
            'format': self.SUPPORTED_FORMATS.get(ext, 'Unknown'),
            'size_bytes': stat.st_size,
            'size_kb': round(stat.st_size / 1024, 2),
            'size_mb': round(stat.st_size / 1024 / 1024, 2),
            'modified_at': stat.st_mtime,
            'is_supported': ext in self.get_supported_formats()
        }


# 便捷函数
def read_file(filepath: str, progress_callback: Callable[[str], None] = None) -> str:
    """便捷函数：读取文件"""
    processor = FileProcessor()
    return processor.read_file(filepath, progress_callback)


def split_text(text: str, max_length: int = 800) -> List[str]:
    """便捷函数：分割文本"""
    processor = FileProcessor()
    return processor.split_text_into_segments(text, max_length)


if __name__ == '__main__':
    # 测试代码
    print("文件处理器模块测试")
    print("=" * 50)

    processor = FileProcessor()

    # 显示支持的格式
    formats = processor.get_supported_formats()
    print(f"\n支持的格式: {list(formats.keys())}")

    print(f"PDF 支持: {PDF_SUPPORT}")
    print(f"EPUB 支持: {EPUB_SUPPORT}")
    print(f"DOCX 支持: {DOCX_SUPPORT}")
    print(f"OCR 支持: {OCR_SUPPORT}")

    # 测试文本分段
    test_text = """
    这是第一段。这段话比较长，包含多个句子。我们要测试分段功能是否正常工作。

    这是第二段。内容不同于第一段。

    这是第三段，也是最后一段。测试结束。
    """

    segments = processor.split_text_into_segments(test_text, max_length=100)
    print(f"\n分段测试 (max_length=100):")
    for i, seg in enumerate(segments):
        print(f"  段 {i+1} ({len(seg)} 字符): {seg[:50]}...")

    # 测试语言检测
    print(f"\n语言检测:")
    print(f"  中文文本: {processor.detect_language('这是一段中文文本')}")
    print(f"  English text: {processor.detect_language('This is an English text')}")
    print(f"  混合文本: {processor.detect_language('Hello 你好 World')}")

    print("\n测试完成!")
